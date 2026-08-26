"""Testes do relatório imprimível da previsão ABC fechada."""
import io
import os
import re
import sys
from pathlib import Path

import pytest
from docx import Document
from PyPDF2 import PdfReader

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from app.services.abc_prediction_report import (
    abc_report_filename,
    build_abc_plain_language_explanation,
    build_abc_report_docx,
    build_abc_report_pdf,
    build_abc_report_pdf_from_docx,
    build_behavior_analyst_narrative,
)
from app.services.abc_closed import build_abc_printable_summary


def _payload() -> dict:
    selected_chain = {
        "cadeia": "Demanda -> Agressão física -> Redirecionamento",
        "antecedente": "Demanda",
        "comportamento": "Agressão física",
        "consequencia": "Redirecionamento",
        "suporte": 8,
        "probabilidade_conjunta": 0.40,
        "probabilidade_comportamento_dado_antecedente": 0.62,
        "probabilidade_consequencia_dada_cadeia_ab": 0.80,
        "c1_leve": 6,
        "c2_intenso": 2,
        "indice_perigo": 0.40,
        "indice_risco": 0.16,
        "funcao_predominante": "Fuga ou esquiva",
    }
    second_chain = {
        "cadeia": "Espera -> Choro -> Atenção social",
        "antecedente": "Espera",
        "comportamento": "Choro",
        "consequencia": "Atenção social",
        "suporte": 3,
        "probabilidade_conjunta": 0.15,
        "probabilidade_comportamento_dado_antecedente": 0.30,
        "probabilidade_consequencia_dada_cadeia_ab": 0.60,
        "c1_leve": 3,
        "c2_intenso": 0,
        "indice_perigo": 0.20,
        "indice_risco": 0.03,
        "funcao_predominante": "Atenção social",
    }
    analysis = {
        "total_registros": 20,
        "c1_leve": 15,
        "c2_intenso": 3,
        "nao_classificados": 2,
        "serie_temporal": [
            {"data": "2026-06-01", "comportamento": "Agressão física", "quantidade": 3},
            {"data": "2026-06-02", "comportamento": "Agressão física", "quantidade": 2},
            {"data": "2026-06-02", "comportamento": "Choro", "quantidade": 1},
        ],
        "cadeias_completas": [selected_chain, second_chain],
    }
    report_records = []
    report_intervals = []
    for index in range(20):
        use_primary = index < 14
        report_records.append(
            {
                "intervalo_id": f"i{index}",
                "data_hora": f"2026-06-{1 + index // 4:02d}T{9 + index % 4:02d}:00:00-03:00",
                "ambiente": "Sala de aula" if index < 15 else "Sala de terapia",
                "antecedente": "Demanda" if use_primary else "Espera",
                "comportamento": "Agressão física" if use_primary else "Choro",
                "consequencia": "Redirecionamento" if use_primary else "Atenção social",
                "classificacao": "C2" if index in {2, 7, 11} else "C1",
                "funcao": "Fuga ou esquiva" if use_primary else "Atenção social",
            }
        )
        report_intervals.append(
            {
                "intervalo_id": f"i{index}",
                "session_id": f"s{index // 4}",
                "status_observacao": "observed",
                "duracao_planejada_minutos": 5,
                "instrumento_versao": "1",
                "eventos_nao_informados": 0,
                "eventos_sem_revisao": 1 if index == 19 else 0,
            }
        )
    report_intervals.extend(
        [
            {"intervalo_id": "i20", "session_id": "s5", "status_observacao": "not_observed", "duracao_planejada_minutos": 5, "instrumento_versao": "1", "eventos_nao_informados": 1},
            {"intervalo_id": "i21", "session_id": "s5", "status_observacao": "invalid", "duracao_planejada_minutos": 5, "instrumento_versao": "1"},
        ]
    )
    temporal_candidates = [
        {
            "origin_behavior_code": "COM_AGRESSAO",
            "from_consequence_code": "CON_REDIRECIONAMENTO",
            "to_antecedent_code": "ANT_DEMANDA",
            "next_behavior_code": "COM_AGRESSAO",
            "delta_seconds": delta,
            "chain_confidence": 0.96,
            "validation_status": "accepted",
            "from_session_id": "s1",
            "to_session_id": "s2",
        }
        for delta in (12, 18, 24)
    ]
    report_summary = build_abc_printable_summary(
        patient_name="Paciente Teste",
        records=report_records,
        intervals=report_intervals,
        temporal_candidates=temporal_candidates,
    )
    report_summary["report_metadata"].update({"period_start": "2026-06-01", "period_end": "2026-06-05"})
    return {
        "patient": "Paciente Teste",
        "environment": "Sala de aula",
        "risk_environment": "Sala de aula",
        "analysis": analysis,
        "analysis_all": {
            **analysis,
            "mapa_risco": [
                {
                    "ambiente": "Sala de aula",
                    "cadeia": selected_chain["cadeia"],
                    "comportamento": "Agressão física",
                    "suporte": 8,
                    "probabilidade_ocorrencia": 0.40,
                    "indice_perigo": 0.40,
                    "classificacao_predominante": "C1",
                }
            ],
        },
        "chain_scope": "selected",
        "selected_chain": selected_chain,
        "prediction": {
            "analysis_mode": "descriptive",
            "comportamento": "Agressão física",
            "estimativa_descritiva": 0.58,
            "probabilidade_prevista": 0.58,
            "numerador": 11,
            "denominador": 19,
            "metodo": "wilson",
            "probabilidade_contexto": 0.55,
            "probabilidade_classificacao": 0.75,
            "probabilidade_funcao": 0.70,
            "intervalo_wilson_inferior": 0.36,
            "intervalo_wilson_superior": 0.76,
            "qualidade_evidencia": "moderada",
        },
        "temporal_data": {
            "candidates": [
                {
                    "origin_behavior_code": "COM_AGRESSAO",
                    "from_consequence_code": "CON_REDIRECIONAMENTO",
                    "to_antecedent_code": "ANT_DEMANDA",
                    "next_behavior_code": "COM_AGRESSAO",
                    "environment": "Sala de aula",
                    "delta_seconds": 12,
                    "validation_status": "accepted",
                    "rejection_reason": None,
                }
            ],
            "stats": [{"insufficient_sample": False}],
        },
        "report_summary": report_summary,
        "ai_result": {"resposta": "Hipótese concorrente: fuga ou esquiva. Confirmar com observação direta."},
    }


def test_explicacao_abc_e_direta_e_sem_nan():
    explicacao = build_abc_plain_language_explanation(_payload())
    texto = " ".join([explicacao["headline"], *explicacao["bullets"]]).lower()

    assert "frequência observada" in texto
    assert "não uma previsão do próximo registro" in texto
    assert "peso configurado de gravidade" in texto
    assert re.search(r"\bnan\b", texto) is None


def test_narrativa_clinica_separa_observacao_hipotese_e_plano_de_mensuracao():
    narrativa = build_behavior_analyst_narrative(_payload())
    texto = " ".join([narrativa["context"], narrativa["objective"], *narrativa["assessment"], *narrativa["hypotheses"], *narrativa["measurement_plan"]])

    assert "exclusivamente registros ABC diretos" in texto
    assert "hipótese concorrente" in texto
    assert "oportunidades em que os mesmos antecedentes ocorreram sem a resposta-alvo" in texto
    assert "idade, diagnóstico" in texto
    assert "constitui diagnóstico" in narrativa["clinical_boundary"]


def test_docx_abc_prioriza_narrativa_e_mantem_formulas_no_apendice():
    conteudo = build_abc_report_docx(_payload())
    doc = Document(io.BytesIO(conteudo))
    texto = "\n".join([*(paragrafo.text for paragrafo in doc.paragraphs), *(cell.text for table in doc.tables for row in table.rows for cell in row.cells)])

    assert "2. Síntese clínica do período" in texto
    assert "Contexto antecedente" in texto
    assert "P(B|A,E)" in texto
    assert "não demonstra" in texto
    assert "Apêndice A — Auditoria dos dados e método" in texto
    assert len(doc.inline_shapes) >= 1


def test_docx_abc_todas_as_cadeias_inclui_guia_dos_indicadores():
    payload = _payload()
    payload["chain_scope"] = "all"
    conteudo = build_abc_report_docx(payload)
    doc = Document(io.BytesIO(conteudo))
    paragrafos = "\n".join(paragrafo.text for paragrafo in doc.paragraphs)
    tabelas = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    texto = f"{paragrafos}\n{tabelas}"

    assert "4. Padrões ABC clinicamente prioritários" in texto
    assert "Demanda -> Agressão física -> Redirecionamento" in texto
    assert "Espera -> Choro -> Atenção social" in texto
    assert "Apêndice A — Auditoria dos dados e método" in texto
    assert "IC de Wilson" in texto
    assert "peso médio" in texto
    assert "não é probabilidade clínica" in texto


def test_docx_abc_cadeia_especifica_mantem_lista_completa_no_apendice():
    conteudo = build_abc_report_docx(_payload())
    doc = Document(io.BytesIO(conteudo))
    tabelas = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "Demanda -> Agressão física -> Redirecionamento" in tabelas
    assert "Espera -> Choro -> Atenção social" in tabelas


def test_pdf_abc_contem_previsao_e_aviso():
    conteudo = build_abc_report_pdf(_payload())
    leitor = PdfReader(io.BytesIO(conteudo))
    texto = "\n".join(pagina.extract_text() or "" for pagina in leitor.pages)
    texto_normalizado = " ".join(texto.split())

    assert conteudo.startswith(b"%PDF")
    assert len(leitor.pages) >= 3
    assert "Relatório descritivo de análise ABC" in texto
    assert "Agressão física" in texto
    assert "Síntese clínica do período" in texto
    assert "Organização temporal dos episódios" in texto
    assert "Gerado em:" in texto
    assert "Apêndice A — Auditoria dos dados e método" in texto
    assert "Padrões ABC clinicamente prioritários" in texto
    assert "Limites, segurança e uso responsável" in texto
    assert "não constitui diagnóstico" in texto_normalizado.lower()


def test_pdf_abc_com_todas_as_cadeias_traz_indicadores_e_cadeias():
    payload = _payload()
    payload["chain_scope"] = "all"
    conteudo = build_abc_report_pdf(payload)
    leitor = PdfReader(io.BytesIO(conteudo))
    texto = "\n".join(pagina.extract_text() or "" for pagina in leitor.pages)

    assert "Padrões ABC clinicamente prioritários" in texto
    assert "Apêndice B — Lista completa das cadeias" in texto
    assert "Demanda -> Agressão física -> Redirecionamento" in texto
    assert "Espera -> Choro -> Atenção social" in texto
    assert "IC de Wilson" in texto
    assert "P(A|E)" in texto


def test_pdf_abc_respeita_anonimizacao_e_secoes_opcionais():
    payload = _payload()
    record = {
        "intervalo_id": "i1",
        "data_hora": "2026-06-01T09:00:00-03:00",
        "ambiente": "Sala de aula",
        "antecedente": "Demanda",
        "comportamento": "Agressão física",
        "consequencia": "Redirecionamento",
        "classificacao": "C1",
        "funcao": "Fuga ou esquiva",
    }
    payload["report_summary"] = build_abc_printable_summary(
        patient_name="Paciente Teste",
        records=[record],
        intervals=[
            {
                "intervalo_id": "i1",
                "session_id": "s1",
                "status_observacao": "observed",
                "duracao_planejada_minutos": 5,
                "instrumento_versao": "1",
            }
        ],
        anonymize_patient=True,
    )
    payload["include_charts"] = False
    payload["include_abc_explanation"] = False
    payload["include_limitations"] = False

    conteudo = build_abc_report_pdf(payload)
    texto = " ".join(
        "\n".join(pagina.extract_text() or "" for pagina in PdfReader(io.BytesIO(conteudo)).pages).split()
    )

    assert "Paciente anonimizado" in texto
    assert "Paciente Teste" not in texto
    assert "O que é o registro ABC fechado?" not in texto
    assert "Limites, segurança e uso responsável" not in texto
    assert "Apêndice A — Auditoria dos dados e método" in texto


def test_css_contem_regras_especificas_para_impressao():
    css = (Path(__file__).parents[1] / "app" / "web" / "assets" / "style.css").read_text(encoding="utf-8")

    assert "@media print" in css
    assert "break-inside: avoid" in css
    assert "display: table-header-group" in css


def test_css_preserva_controle_para_reabrir_sidebar_e_libera_largura():
    css = (Path(__file__).parents[1] / "app" / "web" / "assets" / "style.css").read_text(encoding="utf-8")

    hidden_selectors = re.findall(r"([^{}]+)\{[^{}]*visibility:\s*hidden", css)
    assert not any(re.search(r"(?:^|,)\s*header\s*(?:,|$)", selectors) for selectors in hidden_selectors)
    assert '[data-testid="stHeader"]' in css
    assert 'button[data-testid="stExpandSidebarButton"]' in css
    assert '[data-testid="stSidebar"][aria-expanded="false"]' in css
    assert "flex: 0 0 0 !important" in css
    assert "width: 0 !important" in css


def test_dashboard_expoe_previa_impressao_pdf_e_filtros_do_resumo_abc():
    dashboard = (Path(__file__).parents[1] / "app" / "web" / "dashboard.py").read_text(encoding="utf-8")

    assert "Visualizar resumo" in dashboard
    assert "Imprimir resumo ABC" in dashboard
    assert "Salvar como PDF" in dashboard
    assert "Incluir gráficos" in dashboard
    assert "Apenas revisadas e aceitas" in dashboard
    assert "Ocultar identificação do paciente" in dashboard
    assert "Excel está sendo gerado em segundo plano" in dashboard
    assert "report_docx_bytes = build_abc_report_docx(report_payload)" in dashboard
    assert '"pdf": build_abc_report_pdf_from_docx(report_docx_bytes)' in dashboard


def test_exportacao_fiel_rejeita_conteudo_que_nao_seja_docx():
    with pytest.raises(ValueError, match="Word válido"):
        build_abc_report_pdf_from_docx(b"conteudo invalido")


def test_nome_do_relatorio_abc_e_compativel_com_windows():
    nome = abc_report_filename("João / Teste", "pdf")

    assert nome.endswith(".pdf")
    assert "Joao_Teste" in nome
    assert "/" not in nome
    assert nome.startswith("relatorio_analise_abc_")


def test_resumo_v3_expoe_exposicao_rastreabilidade_e_incerteza_por_sessao():
    payload = _payload()
    summary = payload["report_summary"]
    metadata = summary["report_metadata"]
    quality = summary["data_quality"]
    exposure = summary["exposure_summary"]

    assert metadata["methodology_version"] == "abc-methodology-v3"
    assert len(metadata["dataset_hash"]) == 64
    assert exposure["status"] == "calculable"
    assert exposure["observed_hours"] > 0
    assert "opportunities_without_occurrence" in exposure
    assert summary["top_behavior_session_cluster_interval"]["cluster_count"] >= 2
    assert "invalid_timestamps" in quality
    assert "observer_count" in quality


def test_docx_v3_documenta_exposicao_e_incerteza_sem_poluir_corpo_clinico():
    doc = Document(io.BytesIO(build_abc_report_docx(_payload())))
    text = "\n".join(
        [*(paragraph.text for paragraph in doc.paragraphs), *(cell.text for table in doc.tables for row in table.rows for cell in row.cells)]
    )
    assert "Distribuição da observação" in text
    assert "IC agrupado por sessão" in text
    assert "Indicadores descritivos" in text
