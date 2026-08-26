from __future__ import annotations

import io
import math
import re
import unicodedata
from datetime import date, datetime
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PASTEL_GREEN = "#657f61"
PASTEL_RED = "#bd7d75"
PASTEL_BLUE = "#6f86a4"
PASTEL_GOLD = "#d9c58e"
TEXT = "#2f2a24"
MUTED = "#6f675e"
SURFACE = "#fffaf2"


def behavior_report_filename(patient: str, behavior: str, extension: str) -> str:
    raw = f"relatorio_previsao_{patient}_{behavior}"
    normalized = unicodedata.normalize("NFKD", raw)
    ascii_name = "".join(char for char in normalized if not unicodedata.combining(char))
    safe = re.sub(r"[^A-Za-z0-9_-]+", "_", ascii_name).strip("_")
    return f"{safe[:150] or 'relatorio_previsao_comportamental'}.{extension.lstrip('.')}"


def build_plain_language_explanation(metrics: dict[str, Any]) -> dict[str, Any]:
    probability = _finite(metrics.get("occurrence_probability"), 0.0)
    baseline = _finite(metrics.get("baseline"), 0.0)
    current = _finite(metrics.get("current"), 0.0)
    next_prediction = _finite(metrics.get("next_prediction"), current)
    probability_increase = _finite(metrics.get("probability_increase"), 0.0)
    probability_reduction = _finite(metrics.get("probability_reduction"), 0.0)
    lower = _optional_finite(metrics.get("next_lower"))
    upper = _optional_finite(metrics.get("next_upper"))
    data_points = int(metrics.get("data_points") or 0)
    r2 = _finite(metrics.get("r2"), 0.0)
    risk = str(metrics.get("occurrence_risk") or "indefinido")
    metric_label = str(metrics.get("metric_label") or "medida selecionada").lower()

    expected_sessions = round(probability * 100)
    difference_pp = (probability - baseline) * 100
    if difference_pp >= 5:
        comparison = f"acima do histórico em {abs(difference_pp):.1f} pontos percentuais"
    elif difference_pp <= -5:
        comparison = f"abaixo do histórico em {abs(difference_pp):.1f} pontos percentuais"
    else:
        comparison = "próxima do histórico observado"

    relative_change = 0.0 if current == 0 else (next_prediction - current) / abs(current)
    if next_prediction > current and relative_change >= 0.05:
        trajectory = "aumento"
        trajectory_text = "A projeção mensal aponta aumento da medida selecionada."
    elif next_prediction < current and relative_change <= -0.05:
        trajectory = "redução"
        trajectory_text = "A projeção mensal aponta redução da medida selecionada."
    else:
        trajectory = "estabilidade"
        trajectory_text = "A projeção mensal permanece próxima do nível atual."

    if data_points < 6:
        evidence = "inicial"
        evidence_text = "Há poucos meses observados; leia a projeção com cautela e atualize-a com novos dados."
    elif r2 < 0.30:
        evidence = "variável"
        evidence_text = "Os dados oscilaram e a linha de tendência explica pouco dessa variação."
    elif r2 < 0.70:
        evidence = "moderada"
        evidence_text = "A tendência é parcialmente consistente, mas ainda há variação relevante."
    else:
        evidence = "mais consistente"
        evidence_text = "A trajetória histórica foi relativamente consistente dentro do período analisado."

    interval_text = (
        f"A faixa de 80% para o próximo mês vai de {lower:.2f} a {upper:.2f}."
        if lower is not None and upper is not None
        else "A faixa de incerteza do próximo mês não pôde ser estimada."
    )
    headline = (
        f"A chance estimada de o comportamento aparecer na próxima sessão é de {probability * 100:.1f}% "
        f"(risco {risk})."
    )
    bullets = [
        (
            f"Em 100 sessões semelhantes às observadas, o modelo estimaria cerca de {expected_sessions} "
            "sessões com registro do comportamento. Isso é uma média esperada, não uma garantia."
        ),
        f"Essa chance está {comparison}; o histórico foi de {baseline * 100:.1f}% das sessões.",
        (
            f"Na medida '{metric_label}', o último mês ficou em {current:.2f} e o próximo foi projetado em "
            f"{next_prediction:.2f}. {trajectory_text}"
        ),
        (
            f"O modelo calculou {probability_increase:.1f}% de chance de aumento e "
            f"{probability_reduction:.1f}% de chance de atingir a meta de redução informada."
        ),
        f"{interval_text} {evidence_text}",
    ]
    return {
        "headline": headline,
        "bullets": bullets,
        "comparison": comparison,
        "trajectory": trajectory,
        "evidence": evidence,
        "evidence_text": evidence_text,
        "expected_sessions_per_100": expected_sessions,
    }


def build_behavior_report_docx(payload: dict[str, Any]) -> bytes:
    chart = _build_report_chart(payload)
    explanation = build_plain_language_explanation(payload["metrics"])
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    _configure_docx_styles(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Relatório de previsão comportamental")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(47, 42, 36)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        f"Paciente: {payload['patient']}\n"
        f"Comportamento analisado: {payload['behavior']}\n"
        f"Período: {_date_label(payload['period_start'])} a {_date_label(payload['period_end'])}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Resumo em linguagem simples", level=1)
    p = doc.add_paragraph()
    run = p.add_run(explanation["headline"])
    run.bold = True
    for item in explanation["bullets"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Números principais", level=1)
    _add_docx_metrics_table(doc, payload)
    doc.add_page_break()
    doc.add_picture(chart, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "O gráfico superior compara os valores mensais observados com a trajetória projetada e sua faixa de "
        "incerteza. O gráfico inferior compara a chance para a próxima sessão, o histórico e as probabilidades "
        "de aumento ou redução."
    )

    doc.add_heading("Projeção mensal", level=1)
    _add_docx_forecast_table(doc, payload.get("forecast", pd.DataFrame()))

    doc.add_heading("Como o cálculo foi feito", level=1)
    for item in _methodology_lines(payload):
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Como usar esta informação", level=1)
    for item in _clinical_use_lines():
        doc.add_paragraph(item, style="List Bullet")

    warning = doc.add_paragraph()
    warning_run = warning.add_run(
        "Aviso: este relatório oferece apoio analítico. Ele não confirma função comportamental, não substitui "
        "avaliação funcional ou julgamento clínico e não deve produzir decisões restritivas automáticas."
    )
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(128, 62, 55)
    doc.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}.")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_behavior_report_pdf(payload: dict[str, Any]) -> bytes:
    chart = _build_report_chart(payload)
    explanation = build_plain_language_explanation(payload["metrics"])
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.45 * cm,
        leftMargin=1.45 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
        title="Relatório de previsão comportamental",
        author="Sellas Project",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SellasTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=19, leading=23, textColor=colors.HexColor(TEXT), alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle(name="SellasH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#405f3d"), spaceBefore=8, spaceAfter=7))
    styles.add(ParagraphStyle(name="SellasBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, textColor=colors.HexColor(TEXT), spaceAfter=5))
    styles.add(ParagraphStyle(name="SellasWarning", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9.5, leading=13, textColor=colors.HexColor("#803e37"), backColor=colors.HexColor("#f8e8e3"), borderPadding=7, spaceBefore=8))

    story = [
        Paragraph("Relatório de previsão comportamental", styles["SellasTitle"]),
        Paragraph(
            f"<b>Paciente:</b> {_escape(payload['patient'])}<br/>"
            f"<b>Comportamento:</b> {_escape(payload['behavior'])}<br/>"
            f"<b>Período:</b> {_date_label(payload['period_start'])} a {_date_label(payload['period_end'])}",
            styles["SellasBody"],
        ),
        Paragraph("Resumo em linguagem simples", styles["SellasH1"]),
        Paragraph(f"<b>{_escape(explanation['headline'])}</b>", styles["SellasBody"]),
    ]
    for item in explanation["bullets"]:
        story.append(Paragraph(f"• {_escape(item)}", styles["SellasBody"]))
    story.extend(
        [
            Paragraph("Números principais", styles["SellasH1"]),
            _pdf_metrics_table(payload, styles),
            Spacer(1, 8),
            Image(chart, width=16.2 * cm, height=10.2 * cm),
            Paragraph(
                "O primeiro gráfico mostra a trajetória mensal e a faixa de incerteza. O segundo compara as "
                "probabilidades mais importantes para a leitura clínica.",
                styles["SellasBody"],
            ),
            PageBreak(),
            Paragraph("Projeção mensal", styles["SellasH1"]),
            _pdf_forecast_table(payload.get("forecast", pd.DataFrame()), styles),
            Paragraph("Como o cálculo foi feito", styles["SellasH1"]),
        ]
    )
    for index, item in enumerate(_methodology_lines(payload), start=1):
        story.append(Paragraph(f"{index}. {_escape(item)}", styles["SellasBody"]))
    story.append(Paragraph("Como usar esta informação", styles["SellasH1"]))
    for item in _clinical_use_lines():
        story.append(Paragraph(f"• {_escape(item)}", styles["SellasBody"]))
    story.extend(
        [
            Paragraph(
                "Aviso: este relatório oferece apoio analítico. Ele não confirma função comportamental, não "
                "substitui avaliação funcional ou julgamento clínico e não deve produzir decisões restritivas automáticas.",
                styles["SellasWarning"],
            ),
            Spacer(1, 8),
            Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}.", styles["SellasBody"]),
        ]
    )
    document.build(story)
    return output.getvalue()


def _build_report_chart(payload: dict[str, Any]) -> io.BytesIO:
    observed = payload.get("observed", pd.DataFrame()).copy()
    forecast = payload.get("forecast", pd.DataFrame()).copy()
    metrics = payload["metrics"]
    fig, axes = plt.subplots(2, 1, figsize=(10.8, 7.0), gridspec_kw={"height_ratios": [2.1, 1]})
    fig.patch.set_facecolor(SURFACE)
    for axis in axes:
        axis.set_facecolor(SURFACE)
        axis.grid(axis="y", color="#ded0b8", alpha=0.65, linewidth=0.8)
        axis.spines[["top", "right"]].set_visible(False)
        axis.spines[["left", "bottom"]].set_color("#b9aa90")
        axis.tick_params(colors=TEXT, labelsize=9)

    if not observed.empty:
        observed["month"] = pd.to_datetime(observed["month"], errors="coerce")
        axes[0].plot(observed["month"], observed["valor"], marker="o", linewidth=2.2, color=PASTEL_RED, label="Observado")
    if not forecast.empty:
        forecast["month"] = pd.to_datetime(forecast["month"], errors="coerce")
        axes[0].plot(forecast["month"], forecast["prediction"], marker="o", linewidth=2.2, linestyle="--", color=PASTEL_BLUE, label="Previsto")
        axes[0].fill_between(
            forecast["month"],
            forecast["lower"].astype(float),
            forecast["upper"].astype(float),
            color=PASTEL_BLUE,
            alpha=0.18,
            label="Faixa de 80%",
        )
    axes[0].set_title(f"Trajetória de {payload['behavior']}", loc="left", color=TEXT, fontsize=13, fontweight="bold")
    axes[0].set_ylabel(str(payload.get("y_title") or "Medida mensal"), color=TEXT)
    axes[0].xaxis.set_major_formatter(mdates.DateFormatter("%m/%Y"))
    axes[0].legend(frameon=False, ncol=3, loc="upper left")

    labels = ["Próxima sessão", "Histórico", "Aumentar", "Atingir redução"]
    values = [
        _finite(metrics.get("occurrence_probability")) * 100,
        _finite(metrics.get("baseline")) * 100,
        _finite(metrics.get("probability_increase")),
        _finite(metrics.get("probability_reduction")),
    ]
    bars = axes[1].bar(labels, values, color=[PASTEL_GREEN, PASTEL_GOLD, PASTEL_RED, PASTEL_BLUE], width=0.62)
    axes[1].set_ylim(0, 105)
    axes[1].set_ylabel("Probabilidade (%)", color=TEXT)
    axes[1].set_title("Probabilidades principais", loc="left", color=TEXT, fontsize=12, fontweight="bold")
    for bar, value in zip(bars, values, strict=False):
        axes[1].text(bar.get_x() + bar.get_width() / 2, min(101, value + 2), f"{value:.1f}%", ha="center", va="bottom", fontsize=9, color=TEXT)
    fig.autofmt_xdate(rotation=0)
    fig.tight_layout(pad=1.5)
    output = io.BytesIO()
    fig.savefig(output, format="png", dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    output.seek(0)
    return output


def _configure_docx_styles(doc: Document) -> None:
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor(47, 42, 36)
    doc.styles["Normal"].font.size = Pt(9.5)
    doc.styles["Normal"].paragraph_format.space_after = Pt(3)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.02
    doc.styles["Title"].font.size = Pt(22)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(4)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(64, 95, 61)


def _add_docx_metrics_table(doc: Document, payload: dict[str, Any]) -> None:
    metrics = payload["metrics"]
    rows = [
        ("Chance na próxima sessão", _percent(_finite(metrics.get("occurrence_probability")) * 100)),
        ("Histórico de ocorrência", _percent(_finite(metrics.get("baseline")) * 100)),
        ("Classificação de risco", str(metrics.get("occurrence_risk") or "-")),
        ("Nível atual", _number(metrics.get("current"))),
        ("Previsão do próximo mês", _number(metrics.get("next_prediction"))),
        ("Chance de aumento", _percent(metrics.get("probability_increase"))),
        ("Chance de atingir a redução", _percent(metrics.get("probability_reduction"))),
        ("Qualidade da tendência", str(build_plain_language_explanation(metrics)["evidence"])),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Indicador"
    table.rows[0].cells[1].text = "Resultado"
    for cell in table.rows[0].cells:
        _shade_docx_cell(cell, "DCE8D8")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_docx_forecast_table(doc: Document, forecast: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Mês", "Previsão", "Faixa de 80%", "Leitura"]
    for cell, header in zip(table.rows[0].cells, headers, strict=False):
        cell.text = header
        _shade_docx_cell(cell, "EDE1CE")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    if forecast.empty:
        cells = table.add_row().cells
        cells[0].merge(cells[-1]).text = "Sem projeção mensal disponível."
        return
    for _, row in forecast.head(12).iterrows():
        cells = table.add_row().cells
        prediction = _finite(row.get("prediction"))
        cells[0].text = _month_label(row.get("month"))
        cells[1].text = _number(prediction)
        cells[2].text = f"{_number(row.get('lower'))} a {_number(row.get('upper'))}"
        cells[3].text = "valor mensal projetado"


def _pdf_metrics_table(payload: dict[str, Any], styles) -> Table:
    metrics = payload["metrics"]
    data = [
        ["Indicador", "Resultado", "Indicador", "Resultado"],
        ["Próxima sessão", _percent(_finite(metrics.get("occurrence_probability")) * 100), "Histórico", _percent(_finite(metrics.get("baseline")) * 100)],
        ["Risco", str(metrics.get("occurrence_risk") or "-"), "Nível atual", _number(metrics.get("current"))],
        ["Próximo mês", _number(metrics.get("next_prediction")), "Chance de aumento", _percent(metrics.get("probability_increase"))],
        ["Atingir redução", _percent(metrics.get("probability_reduction")), "Qualidade", build_plain_language_explanation(metrics)["evidence"]],
    ]
    wrapped = [[Paragraph(_escape(str(value)), styles["SellasBody"]) for value in row] for row in data]
    table = Table(wrapped, colWidths=[4.0 * cm, 2.5 * cm, 4.0 * cm, 2.7 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dce8d8")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9b897")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def _pdf_forecast_table(forecast: pd.DataFrame, styles) -> Table:
    data = [["Mês", "Previsão", "Faixa de 80%"]]
    if forecast.empty:
        data.append(["-", "Sem projeção", "-"])
    else:
        for _, row in forecast.head(12).iterrows():
            data.append(
                [
                    _month_label(row.get("month")),
                    _number(row.get("prediction")),
                    f"{_number(row.get('lower'))} a {_number(row.get('upper'))}",
                ]
            )
    wrapped = [[Paragraph(_escape(str(value)), styles["SellasBody"]) for value in row] for row in data]
    table = Table(wrapped, colWidths=[4.2 * cm, 4.2 * cm, 6.3 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#ede1ce")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9b897")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table


def _methodology_lines(payload: dict[str, Any]) -> list[str]:
    metrics = payload["metrics"]
    reduction_target = int(metrics.get("reduction_target") or 0)
    return [
        "Os registros foram agrupados por mês usando a medida selecionada.",
        "A trajetória mensal foi ajustada na escala logarítmica: log(valor + 1) = intercepto + beta x mês.",
        "A previsão retorna à escala original por exp(previsão logarítmica) - 1, impedindo valores negativos.",
        "A chance da próxima sessão combina histórico, janela recente, média exponencial e direção da tendência.",
        f"A chance de redução considera a meta informada de {reduction_target}% em relação ao último mês.",
        f"Ajuste da tendência: beta = {_finite(metrics.get('beta')):+.3f}; R² = {_finite(metrics.get('r2')):.2f}; erro sigma = {_finite(metrics.get('sigma')):.2f}.",
    ]


def _clinical_use_lines() -> list[str]:
    return [
        "Compare o resultado com observação direta, oportunidades de ocorrência, ambiente e mudanças de rotina.",
        "Atualize o relatório quando entrarem novos dados; probabilidades podem mudar.",
        "Use a faixa de incerteza e a qualidade da tendência, não apenas um percentual isolado.",
        "Investigue hipóteses funcionais por avaliação apropriada; a previsão sozinha não identifica função.",
    ]


def _shade_docx_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _finite(value: Any, default: float = 0.0) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return default
    return number if math.isfinite(number) else default


def _optional_finite(value: Any) -> float | None:
    if value is None:
        return None
    number = _finite(value, float("nan"))
    return number if math.isfinite(number) else None


def _number(value: Any) -> str:
    return f"{_finite(value):.2f}"


def _percent(value: Any) -> str:
    return f"{_finite(value):.1f}%"


def _month_label(value: Any) -> str:
    timestamp = pd.to_datetime(value, errors="coerce")
    return "-" if pd.isna(timestamp) else timestamp.strftime("%m/%Y")


def _date_label(value: date | datetime | str) -> str:
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    parsed = pd.to_datetime(value, errors="coerce")
    return str(value) if pd.isna(parsed) else parsed.strftime("%d/%m/%Y")


def _escape(value: Any) -> str:
    text = str(value)
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
