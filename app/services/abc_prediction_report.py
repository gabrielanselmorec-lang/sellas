from __future__ import annotations

import io
import math
import os
import re
import tempfile
import textwrap
import threading
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.services.abc_methodology import DEFAULT_CLINICAL_DISCLAIMER, METHODOLOGY_VERSION


TEXT = "#2f2a24"
GREEN = "#657f61"
RED = "#bd7d75"
BLUE = "#6f86a4"
GOLD = "#d9c58e"
SURFACE = "#fffaf2"

_WORD_PDF_EXPORT_LOCK = threading.Lock()


def abc_report_filename(patient: str, extension: str) -> str:
    normalized = unicodedata.normalize("NFKD", f"relatorio_analise_abc_{patient}")
    ascii_name = "".join(char for char in normalized if not unicodedata.combining(char))
    safe = re.sub(r"[^A-Za-z0-9_-]+", "_", ascii_name).strip("_")
    return f"{safe[:150] or 'relatorio_analise_abc'}.{extension.lstrip('.')}"


def build_abc_plain_language_explanation(payload: dict[str, Any]) -> dict[str, Any]:
    analysis = payload.get("analysis") or {}
    prediction = payload.get("prediction") or {}
    chain = payload.get("selected_chain") or {}
    total = int(analysis.get("total_registros") or 0)
    c1 = int(analysis.get("c1_leve") or 0)
    c2 = int(analysis.get("c2_intenso") or 0)
    unclassified = int(analysis.get("nao_classificados") or 0)
    mode = str(prediction.get("analysis_mode") or analysis.get("analysis_mode") or "descriptive")
    estimate_value = prediction.get("estimativa_descritiva", prediction.get("probabilidade_prevista"))
    estimate = _finite(estimate_value) if estimate_value is not None else None
    lower_value = prediction.get("intervalo_wilson_inferior")
    upper_value = prediction.get("intervalo_wilson_superior")
    lower = _finite(lower_value) if lower_value is not None else None
    upper = _finite(upper_value) if upper_value is not None else None
    target_behavior = str(prediction.get("comportamento") or "comportamento selecionado")
    evidence = str(prediction.get("qualidade_evidencia") or "inicial")
    environment = str(payload.get("environment") or "Todos os ambientes")
    chain_scope = str(payload.get("chain_scope") or "selected")
    all_chains = (payload.get("analysis") or {}).get("cadeias_completas") or []

    if mode == "predictive":
        headline = f"O modelo validado estimou '{target_behavior}' conforme o alvo e o horizonte documentados."
    elif estimate is None:
        headline = f"Não há registros semelhantes suficientes para estimar a frequência observada de '{target_behavior}'."
    else:
        k = int(prediction.get("numerador") or prediction.get("sucessos_contexto") or 0)
        n = int(prediction.get("denominador") or prediction.get("amostra_contexto") or 0)
        headline = (
            f"Entre os episódios ABC registrados no contexto selecionado, '{target_behavior}' apareceu em "
            f"{k} de {n} registros ({estimate * 100:.1f}%)."
            if n > 0
            else f"A frequência observada informada para '{target_behavior}' foi de {estimate * 100:.1f}%; o contrato legado não forneceu k e n."
        )
    bullets = [
        (
            "Esta é uma frequência observada entre episódios registrados, não uma previsão do próximo registro nem "
            "uma estimativa de risco absoluto. A base não contém todas as oportunidades observáveis sem comportamento."
            if mode != "predictive"
            else "A seção preditiva informa períodos de treino e teste, baseline, métricas e calibração fora da amostra."
        ),
        (
            f"O intervalo reproduzível vai de {lower * 100:.1f}% a {upper * 100:.1f}%; qualidade '{evidence}'."
            if lower is not None and upper is not None
            else f"Não há intervalo calculável no recorte; qualidade '{evidence}'."
        ),
        (
            f"No recorte '{environment}', há {c1} registros C1 - leves, {c2} C2 - intensos e "
            f"{unclassified} ainda sem classificação. Frequência e peso configurado de gravidade são apresentados separadamente."
        ),
    ]
    if chain:
        chain_probability = _finite(chain.get("probabilidade_conjunta"))
        severity_value = chain.get("peso_medio_gravidade", chain.get("indice_perigo"))
        danger = _finite(severity_value) if severity_value is not None else None
        risk = _finite(chain.get("indice_risco"), chain_probability * danger)
        if chain_scope == "all":
            bullets.append(
                f"O relatório inclui todas as {len(all_chains)} cadeias do recorte. A cadeia mais frequente "
                f"'{chain.get('cadeia', '-')}' é usada como referência visual: ocorreu em {chain_probability * 100:.1f}% "
                f"dos registros, com peso médio configurado de gravidade {_pct_or_na(danger)} e índice exploratório {_pct_or_na(risk)}."
            )
        else:
            bullets.append(
                f"A cadeia selecionada '{chain.get('cadeia', '-')}' ocorreu em {chain_probability * 100:.1f}% dos "
                f"registros do recorte, com peso médio configurado de gravidade {_pct_or_na(danger)} e índice exploratório {_pct_or_na(risk)}."
            )
        bullets.append(
            f"A hipótese funcional mais frequentemente registrada foi '{chain.get('funcao_predominante') or 'não informada'}'. "
            "Ela não confirma função comportamental e não substitui avaliação funcional."
        )
    report_summary = payload.get("report_summary") if isinstance(payload.get("report_summary"), dict) else {}
    exposure = report_summary.get("exposure_summary") or (report_summary.get("observation_summary") or {}).get("exposure") or {}
    if exposure.get("status") == "calculable":
        bullets.append(
            f"A exposição válida totaliza {_finite(exposure.get('observed_hours')):.2f} hora(s), com "
            f"{_number_or_na(exposure.get('occurrences_per_hour'))} episódio(s) por hora observada e "
            f"{int(exposure.get('opportunities_without_occurrence') or 0)} oportunidade(s) observada(s) sem episódio."
        )
    else:
        bullets.append("A exposição temporal não está disponível; contagens por data não devem ser interpretadas como taxas.")
    return {"headline": headline, "bullets": bullets}


def build_behavior_analyst_narrative(payload: dict[str, Any]) -> dict[str, Any]:
    """Transforma os indicadores do relatório em uma formulação clínica auditável.

    A narrativa mantém a distinção entre observação, hipótese e decisão clínica. Dados
    biopsicossociais não presentes na fonte nunca são completados por inferência.
    """
    summary = _normalized_report_summary(payload)
    analysis = payload.get("analysis") or {}
    observation = summary.get("observation_summary") or {}
    frequent = summary.get("most_frequent") or {}
    exposure = summary.get("exposure_summary") or observation.get("exposure") or {}
    total = int(observation.get("total_records") or analysis.get("total_registros") or 0)
    hours = _finite(observation.get("observed_hours") or exposure.get("observed_hours"))
    rate = observation.get("occurrences_per_observed_hour", exposure.get("occurrences_per_hour"))

    context = (
        "Esta análise utiliza exclusivamente registros ABC diretos do período e dos ambientes identificados abaixo. "
        "A fonte consultada não informa idade, diagnóstico, repertórios de comunicação e autonomia, condições de saúde, "
        "prioridades do paciente ou relatos de familiares e da escola. Esses dados devem ser integrados pelo analista "
        "responsável antes de qualquer decisão de intervenção."
    )

    objective_parts = [
        f"Foram incluídos {total} registros ABC no período de {_summary_period_label(summary, payload)}."
    ]
    if hours > 0:
        objective_parts.append(f"A cobertura observacional válida totalizou {hours:.1f} horas.")
    if rate is not None:
        objective_parts.append(
            f"Nesse recorte, foram registrados {_decimal_pt(rate, 1)} episódios por hora observada; "
            "essa taxa descreve a amostra e não estima risco futuro."
        )
    top_behavior = frequent.get("behavior") or {}
    top_antecedent = frequent.get("antecedent") or {}
    top_consequence = frequent.get("consequence") or {}
    if top_behavior.get("nome"):
        objective_parts.append(
            f"A resposta mais registrada foi '{top_behavior.get('nome')}' ({int(top_behavior.get('quantidade') or 0)} ocorrências)."
        )
    if top_antecedent.get("nome") and top_consequence.get("nome"):
        objective_parts.append(
            f"Considerados separadamente, o antecedente mais frequente foi '{top_antecedent.get('nome')}' e a consequência "
            f"imediata mais frequente foi '{top_consequence.get('nome')}'."
        )

    assessment = []
    top_chains = sorted(
        analysis.get("cadeias_completas") or [], key=lambda item: int(item.get("suporte") or 0), reverse=True
    )
    for item in top_chains[:3]:
        support = int(item.get("suporte") or 0)
        assessment.append(
            f"Em {support} registros, observou-se a sequência '{item.get('cadeia') or '-'}'. "
            "A repetição torna esse arranjo prioritário para observação dirigida, mas não demonstra que o antecedente causou "
            "a resposta ou que a consequência a reforçou."
        )
    ambiguous_labels = {
        str(value).strip().casefold()
        for item in top_chains
        for value in (item.get("antecedente"), item.get("consequencia"))
        if value
    }
    if '"do nada"' in ambiguous_labels or "do nada" in ambiguous_labels:
        assessment.append(
            "O rótulo 'do nada' aparece nos registros, mas não constitui uma descrição operacional de antecedente. "
            "Antes de interpretá-lo, a equipe deve registrar atividade em curso, demanda, transição, estímulos presentes, "
            "interações sociais e mudanças ambientais imediatamente anteriores."
        )
    if not assessment:
        assessment.append("Não houve cadeia com suporte suficiente para uma formulação descritiva no corpo principal.")

    functions = sorted(
        analysis.get("por_funcao") or [], key=lambda item: int(item.get("quantidade") or 0), reverse=True
    )
    if not functions:
        selected = payload.get("selected_chain") or {}
        if selected.get("funcao_predominante"):
            functions = [{"funcao": selected.get("funcao_predominante"), "quantidade": selected.get("suporte") or 0}]
    hypotheses = []
    for item in functions[:4]:
        name = str(item.get("funcao") or "Não identificada")
        count = int(item.get("quantidade") or 0)
        hypotheses.append(
            f"{name}: rótulo funcional registrado em {count} episódios. Deve ser tratado como hipótese concorrente, "
            "a ser contrastada com antecedentes, consequências, operações motivadoras e oportunidades sem ocorrência."
        )
    if not hypotheses:
        hypotheses.append("Nenhuma hipótese funcional foi registrada de modo utilizável neste recorte.")

    measurement_plan = [
        "Integrar idade, repertórios de comunicação e autorregulação, condições de saúde, preferências, objetivos do plano e relatos do paciente, da família e da escola.",
        "Revisar e alinhar as definições operacionais das respostas prioritárias, registrando início, término, duração, intensidade observável e impacto sobre segurança e participação.",
        "Substituir, nas coletas futuras, rótulos amplos como 'do nada' e 'comportamento bloqueado' por descrições observáveis do evento, da ação do mediador e da mudança ambiental.",
        "Registrar oportunidades em que os mesmos antecedentes ocorreram sem a resposta-alvo; isso permite comparar episódios com e sem ocorrência e testar as hipóteses concorrentes.",
        "Para as cadeias prioritárias, medir latência após o antecedente, duração da resposta, consequência efetivamente produzida e retorno à atividade ou ao acesso socialmente relevante.",
        "Planejar qualquer apoio antecedente, ensino de comunicação funcional, escolha, previsibilidade ou reforçamento diferencial somente após revisão do analista responsável, com critérios de fidelidade, segurança e benefício para o paciente.",
    ]

    return {
        "context": context,
        "objective": " ".join(objective_parts),
        "assessment": assessment,
        "hypotheses": hypotheses,
        "measurement_plan": measurement_plan,
        "clinical_boundary": (
            "Este documento é uma análise descritiva de registros ABC. Não constitui diagnóstico, análise funcional "
            "experimental, plano de intervenção comportamental ou autorização para procedimentos restritivos."
        ),
    }


def build_abc_report_docx(payload: dict[str, Any]) -> bytes:
    summary = _normalized_report_summary(payload)
    narrative = build_behavior_analyst_narrative(payload)
    include_charts = bool(payload.get("include_charts", True))
    chain_chart = _build_chain_frequency_chart(payload, summary) if include_charts else None
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    _configure_docx_styles(doc)
    _configure_docx_header_footer(section)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Relatório descritivo de análise ABC")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(23)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    subtitle = doc.add_paragraph("Formulação clínica orientada ao analista do comportamento, com trilha técnica de auditoria")
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_heading("1. Identificação e contexto disponível", level=1)
    _add_docx_table(
        doc,
        ["Campo", "Valor"],
        _clinical_identification_rows(payload),
        [1.9, 4.6],
    )
    doc.add_paragraph(narrative["context"])

    doc.add_heading("2. Síntese clínica do período", level=1)
    doc.add_heading("Contexto relatado", level=2)
    doc.add_paragraph(
        "Não há, nesta fonte, relatos subjetivos do paciente, da família ou da equipe escolar. "
        "A ausência é registrada para evitar que inferências sejam apresentadas como fatos clínicos."
    )
    doc.add_heading("Observações diretas", level=2)
    doc.add_paragraph(narrative["objective"])
    doc.add_heading("Formulação analítico-comportamental", level=2)
    for item in narrative["assessment"]:
        doc.add_paragraph(item)
    doc.add_heading("Próximos passos de análise e mensuração", level=2)
    for item in narrative["measurement_plan"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Distribuição da observação", level=1)
    doc.add_paragraph(
        "A distribuição por ambiente ajuda a delimitar onde a amostra foi construída. Diferenças de taxa devem ser "
        "interpretadas junto da duração, das atividades e das oportunidades observadas em cada contexto."
    )
    _add_docx_table(doc, ["Ambiente", "Horas", "Episódios", "Episódios/h"], _exposure_rows(summary), [3.2, 1.0, 1.1, 1.2])

    doc.add_page_break()
    doc.add_heading("4. Padrões ABC clinicamente prioritários", level=1)
    doc.add_paragraph(
        "As sequências abaixo foram priorizadas pela repetição no recorte. A leitura conserva a ordem observada "
        "— contexto antecedente, resposta e mudança ambiental imediata — sem atribuir causalidade."
    )
    _add_docx_table(
        doc,
        ["Contexto antecedente", "Resposta observada", "Evento posterior", "n", "Leitura"],
        _clinical_chain_rows(payload),
        [1.45, 1.45, 1.45, 0.7, 1.45],
    )
    if include_charts and chain_chart is not None:
        doc.add_picture(chain_chart, width=Inches(6.45))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Frequência das cadeias priorizadas; o gráfico não representa probabilidade de risco ou função confirmada.")

    doc.add_heading("5. Hipóteses funcionais concorrentes", level=1)
    doc.add_paragraph(
        "Os rótulos funcionais registrados são tratados como hipóteses concorrentes. Sua frequência indica onde aprofundar "
        "a avaliação; não confirma a função operante de uma resposta."
    )
    for item in narrative["hypotheses"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "A confirmação requer contraste entre condições relevantes, oportunidades sem ocorrência, observação direta e julgamento profissional."
    )

    doc.add_heading("6. Organização temporal dos episódios", level=1)
    doc.add_paragraph(
        "As transições agrupadas descrevem episódios próximos no tempo. Elas ajudam a formular perguntas sobre recuperação, "
        "retorno à atividade e efeitos do manejo, mas não demonstram uma contingência de reforçamento."
    )
    _add_docx_table(
        doc,
        ["Transição agrupada", "n", "Sess.", "Dias", "Med. (s)", "IQR (s)", "Status"],
        _temporal_rows_v2(payload),
        [2.65, 0.4, 0.65, 0.5, 0.75, 0.65, 0.9],
    )
    doc.add_paragraph(
        "Nos registros legados, 'Ignorado' é preservado como valor oficial e significa que não houve mudança ambiental específica registrada; "
        "não significa omissão de cuidado."
    )

    doc.add_heading("7. Limites, segurança e uso responsável", level=1)
    warning = doc.add_paragraph()
    warning_run = warning.add_run(narrative["clinical_boundary"])
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(128, 62, 55)
    doc.add_paragraph(
        "A linguagem descreve respostas observáveis e barreiras de participação sem atribuir culpa, intenção ou traço pessoal. "
        "O relatório não recomenda automaticamente extinção, bloqueio, restrição, contenção ou qualquer procedimento clínico."
    )
    doc.add_paragraph(
        "Os achados dependem da completude dos registros, das definições operacionais e da representatividade do período. "
        "A generalização para outros ambientes, observadores e períodos permanece desconhecida."
    )

    doc.add_heading("Apêndice A — Auditoria dos dados e método", level=1)
    doc.add_paragraph(
        "Esta seção preserva a rastreabilidade técnica sem interromper a formulação clínica. Indicadores de banco de dados, "
        "estimativas e fórmulas destinam-se à revisão interna e não devem ser lidos isoladamente como conclusões clínicas."
    )
    _add_docx_table(doc, ["Indicador", "Resultado"], _quality_rows(summary), [3.9, 2.6])
    doc.add_heading("Indicadores descritivos", level=2)
    _add_docx_table(doc, ["Indicador", "Resultado"], _descriptive_metric_rows(payload), [3.9, 2.6])
    doc.add_heading("Método de cálculo", level=2)
    for item in _methodology_lines(payload):
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("Apêndice B — Lista completa das cadeias", level=1)
    doc.add_paragraph(
        "Os rótulos legados permanecem reproduzidos para rastreabilidade. Nesta leitura, 'Ignorado' designa ausência de mudança ambiental específica registrada."
    )
    all_chain_rows = _all_chain_rows(payload)
    _add_docx_table(
        doc,
        ["Cadeia", "k", "n", "Freq.", "IC inf.", "IC sup.", "NC"],
        all_chain_rows,
        [2.6, 0.45, 0.45, 0.8, 0.75, 0.75, 0.7],
        prevent_row_split=False,
    )
    doc.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}; metodologia {METHODOLOGY_VERSION}.")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_abc_report_pdf_from_docx(docx_bytes: bytes) -> bytes:
    """Exporta para PDF exatamente o DOCX entregue ao usuário.

    A exportação pelo Microsoft Word evita que DOCX e PDF sejam diagramados
    por mecanismos diferentes. O bloqueio impede duas instâncias concorrentes
    do Word no mesmo processo do Streamlit.
    """
    document_bytes = bytes(docx_bytes) if isinstance(docx_bytes, (bytes, bytearray)) else b""
    if not document_bytes.startswith(b"PK"):
        raise ValueError("O conteúdo informado não é um arquivo Word válido.")
    if os.name != "nt":
        raise RuntimeError("A exportação fiel para PDF requer Microsoft Word no Windows.")

    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError("A exportação fiel para PDF requer Microsoft Word e pywin32.") from exc

    with _WORD_PDF_EXPORT_LOCK, tempfile.TemporaryDirectory(prefix="sellas_abc_pdf_") as temp_dir:
        docx_path = Path(temp_dir) / "relatorio_abc.docx"
        pdf_path = Path(temp_dir) / "relatorio_abc.pdf"
        docx_path.write_bytes(document_bytes)

        pythoncom.CoInitialize()
        word = None
        document = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            word.Options.UpdateLinksAtOpen = False
            document = word.Documents.Open(
                str(docx_path.resolve()),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
                OpenAndRepair=True,
                NoEncodingDialog=True,
            )
            document.ExportAsFixedFormat(
                OutputFileName=str(pdf_path.resolve()),
                ExportFormat=17,
                OpenAfterExport=False,
                OptimizeFor=0,
                Range=0,
                Item=0,
                IncludeDocProps=True,
                KeepIRM=True,
                CreateBookmarks=1,
                DocStructureTags=True,
                BitmapMissingFonts=True,
                UseISO19005_1=False,
            )
            if not pdf_path.exists():
                raise RuntimeError("O Microsoft Word não produziu o arquivo PDF.")
            pdf_bytes = pdf_path.read_bytes()
            if not pdf_bytes.startswith(b"%PDF"):
                raise RuntimeError("O arquivo exportado pelo Microsoft Word não é um PDF válido.")
            return pdf_bytes
        except Exception as exc:
            if isinstance(exc, RuntimeError):
                raise
            raise RuntimeError(f"Não foi possível exportar o Word para PDF: {exc}") from exc
        finally:
            if document is not None:
                document.Close(False)
            if word is not None:
                word.Quit()
            pythoncom.CoUninitialize()


def _build_abc_report_pdf_legacy(payload: dict[str, Any]) -> bytes:
    explanation = build_abc_plain_language_explanation(payload)
    summary = _normalized_report_summary(payload)
    include_charts = bool(payload.get("include_charts", True))
    include_explanation = bool(payload.get("include_abc_explanation", True))
    include_limitations = bool(payload.get("include_limitations", True))
    overview_chart = _build_overview_chart(payload) if include_charts else None
    probability_chart = _build_probability_chart(payload) if include_charts else None
    distribution_chart = _build_summary_distribution_chart(summary) if include_charts else None
    association_chart = _build_association_heatmap(summary) if include_charts else None
    chain_chart = _build_chain_frequency_chart(payload, summary) if include_charts else None
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.35 * cm,
        leftMargin=1.35 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.25 * cm,
        title="Relatório de análise ABC fechada",
        author="Sellas Project",
    )
    styles = _pdf_styles()
    metadata = summary.get("report_metadata") or {}
    patient = summary.get("patient") or {}
    observation = summary.get("observation_summary") or {}
    quality = summary.get("data_quality") or {}
    story = [
        Paragraph("Relatório de análise ABC fechada", styles["SellasTitle"]),
        Paragraph(
            f"<b>Paciente:</b> {_escape(patient.get('display_name') or payload.get('patient', '-'))}<br/>"
            f"<b>Ambiente:</b> {_escape(payload.get('environment') or 'Todos os ambientes')}<br/>"
            f"<b>Período:</b> {_escape(_summary_period_label(summary, payload))}<br/>"
            f"<b>Responsável:</b> {_escape(metadata.get('generated_by') or 'Usuário local')}<br/>"
            f"<b>Instrumento:</b> ABC fechado, versão {_escape(metadata.get('instrument_version') or 'não disponível')}<br/>"
            f"<b>Lógica analítica:</b> {_escape(metadata.get('logic_version') or 'abc-print-summary-v1')}<br/>"
            f"<b>Gerado em:</b> {_escape(_generated_at_label(metadata.get('generated_at')))}",
            styles["SellasBody"],
        ),
        Paragraph("Resumo descritivo dos registros", styles["SellasH1"]),
        Paragraph(_escape(summary.get("descriptive_summary") or explanation["headline"]), styles["SellasBody"]),
        Paragraph("Informações gerais", styles["SellasH1"]),
        _pdf_table(
            ["Indicador", "Resultado"],
            _clinical_summary_rows(summary, payload),
            [9.5 * cm, 6.3 * cm],
            styles,
        ),
    ]
    story.append(Paragraph(f"<b>Qualidade geral:</b> {_escape(str(quality.get('status_scope') or quality.get('status') or 'insuficiente').capitalize())}", styles["SellasBody"]))
    story.append(Paragraph(_escape(summary.get("clinical_disclaimer") or ""), styles["SellasWarning"]))

    if include_explanation:
        story.extend(
            [
                PageBreak(),
                Paragraph("O que é o registro ABC fechado?", styles["SellasH1"]),
                Paragraph(
                    "O registro ABC fechado é uma forma estruturada de observar e registrar eventos relacionados a um comportamento. "
                    "O aplicador seleciona opções previamente definidas para antecedente, comportamento e consequência.",
                    styles["SellasBody"],
                ),
                _pdf_table(
                    ["Componente", "Definição e exemplos"],
                    _abc_component_rows(),
                    [3.2 * cm, 12.6 * cm],
                    styles,
                ),
                Spacer(1, 8),
                Paragraph("<b>Antecedente -&gt; Comportamento -&gt; Consequência</b>", styles["SellasFlow"]),
                Paragraph(
                    "O registro descreve o que foi observado. Ele não confirma sozinho a função do comportamento e não substitui avaliação funcional realizada por profissional qualificado.",
                    styles["SellasWarning"],
                ),
                Paragraph("Como os dados são registrados?", styles["SellasH1"]),
                Paragraph(
                    "A sessão é dividida em intervalos, normalmente de cinco minutos. Os cálculos usam intervalos distintos e efetivamente observados. "
                    "Mais de um evento pode aparecer no mesmo intervalo, e intervalos sem observação válida ficam fora dos denominadores de ocorrência.",
                    styles["SellasBody"],
                ),
                _pdf_table(
                    ["Estado", "Significado"],
                    _observation_state_rows(),
                    [3.2 * cm, 12.6 * cm],
                    styles,
                ),
                Paragraph("Valores nulos ou não observados nunca são convertidos automaticamente em 'não ocorreu'.", styles["SellasWarning"]),
            ]
        )

    story.extend(
        [
            PageBreak(),
            Paragraph("Qualidade e cobertura dos dados", styles["SellasH1"]),
            _pdf_table(
                ["Indicador", "Resultado"],
                _quality_rows(summary),
                [9.3 * cm, 6.5 * cm],
                styles,
            ),
            Paragraph("Exposição observada por ambiente", styles["SellasH1"]),
            _pdf_table(
                ["Ambiente", "Horas", "Episódios", "Episódios/h"],
                _exposure_rows(summary),
                [7.4 * cm, 2.6 * cm, 2.8 * cm, 3.0 * cm],
                styles,
            ),
        ]
    )
    warnings = quality.get("warnings") or []
    if warnings:
        story.append(Paragraph("Alertas de qualidade", styles["SellasH1"]))
        for warning in warnings:
            story.append(Paragraph(f"• {_escape(warning)}", styles["SellasBody"]))
    else:
        story.append(Paragraph("Nenhum alerta estrutural foi identificado no recorte.", styles["SellasBody"]))

    if include_charts:
        story.extend(
            [
                PageBreak(),
                Paragraph("Frequências resumidas", styles["SellasH1"]),
                Image(distribution_chart, width=17.2 * cm, height=11.0 * cm),
                Paragraph(
                    "As barras mostram contagens observadas. Os rótulos mantêm a leitura possível mesmo sem diferenciação por cor.",
                    styles["SellasBody"],
                ),
                PageBreak(),
                Paragraph("Linha temporal e mapa de frequência por gravidade configurada", styles["SellasH1"]),
                Image(overview_chart, width=17.7 * cm, height=10.8 * cm),
                Paragraph(
                    "Frequência e peso configurado de gravidade são dimensões diferentes. Os pesos não são probabilidades clínicas nem escala validada.",
                    styles["SellasBody"],
                ),
            ]
        )

    story.extend(
        [
            PageBreak(),
            Paragraph("Principais associações antecedente-comportamento", styles["SellasH1"]),
            Paragraph(
                "As linhas abaixo apresentam associações observadas e hipóteses para análise clínica. Elas não identificam gatilhos ou causas confirmadas.",
                styles["SellasBody"],
            ),
            _pdf_table(
                ["Antecedente", "Comportamento", "Exposições", "Conjuntas", "P(B|A)", "Base", "Dif. risco", "Lift", "Qualidade"],
                _association_table_rows(summary, "antecedent_behavior_associations", "antecedente", "comportamento"),
                [3.3 * cm, 3.0 * cm, 1.45 * cm, 1.35 * cm, 1.45 * cm, 1.25 * cm, 1.45 * cm, 1.0 * cm, 2.25 * cm],
                styles,
            ),
            Paragraph("Principais associações comportamento-consequência", styles["SellasH1"]),
            _pdf_table(
                ["Comportamento", "Consequência", "Exposições", "Conjuntas", "P(C|B)", "Base", "Dif. risco", "Lift", "Qualidade"],
                _association_table_rows(summary, "behavior_consequence_associations", "comportamento", "consequencia"),
                [3.3 * cm, 3.0 * cm, 1.45 * cm, 1.35 * cm, 1.45 * cm, 1.25 * cm, 1.45 * cm, 1.0 * cm, 2.25 * cm],
                styles,
            ),
        ]
    )
    if include_charts:
        story.extend(
            [
                PageBreak(),
                Paragraph("Matrizes das associações observadas", styles["SellasH1"]),
                Image(association_chart, width=17.2 * cm, height=10.8 * cm),
                Paragraph("Cada célula mostra uma probabilidade condicional observada no recorte.", styles["SellasBody"]),
            ]
        )

    story.extend(
        [
            PageBreak(),
            Paragraph(_chain_section_heading(payload), styles["SellasH1"]),
            _pdf_table(
                ["Cadeia", "n", "P(A-B-C)", "C1", "C2", "Peso médio"],
                _top_chain_rows(payload),
                [8.1 * cm, 0.8 * cm, 2.0 * cm, 0.8 * cm, 0.8 * cm, 1.6 * cm],
                styles,
            ),
        ]
    )
    if include_charts:
        story.extend([Spacer(1, 8), Image(chain_chart, width=17.0 * cm, height=7.0 * cm)])
    story.extend(
        [
            Paragraph("Como são identificadas cadeias comportamentais?", styles["SellasH1"]),
            Paragraph(
                "Uma sessão pode conter uma sequência contínua na qual a consequência de um episódio modifica o ambiente e passa a compor o antecedente do episódio seguinte.",
                styles["SellasBody"],
            ),
            Paragraph("<b>A1 -&gt; B1 -&gt; C1 =&gt; A2 -&gt; B2 -&gt; C2</b>", styles["SellasFlow"]),
            Paragraph("Exemplo de leitura temporal", styles["SellasH1"]),
            Paragraph("1. Uma demanda foi apresentada antes de um protesto vocal.", styles["SellasBody"]),
            Paragraph("2. Após o protesto, a demanda foi temporariamente retirada.", styles["SellasBody"]),
            Paragraph("3. A pausa resultante passou a compor o contexto antecedente do episódio seguinte.", styles["SellasBody"]),
            Paragraph("4. Quando a demanda retornou, outro comportamento pôde ser observado.", styles["SellasBody"]),
            Paragraph("5. A aplicação conecta os episódios somente quando a ordem temporal, a regra de transição versionada e a revisão registrada são atendidas.", styles["SellasBody"]),
            Paragraph(
                "A identificação respeita a ordem temporal. Uma cadeia detectada é um padrão temporal descritivo e não demonstra, isoladamente, que uma consequência reforçou ou causou o comportamento seguinte.",
                styles["SellasWarning"],
            ),
            Paragraph("Cadeias temporais incluídas", styles["SellasH1"]),
            _pdf_table(
                ["Cadeia", "Repet.", "Sess.", "Med.", "Mín.", "Máx.", "Estável", "Conf.", "Status"],
                _temporal_summary_rows(summary),
                [5.5 * cm, 1.0 * cm, 1.25 * cm, 1.25 * cm, 1.05 * cm, 1.05 * cm, 1.4 * cm, 1.45 * cm, 2.4 * cm],
                styles,
            ),
            PageBreak(),
            Paragraph("Estimativa descritiva e indicadores", styles["SellasH1"]),
        ]
    )
    if include_charts:
        story.append(Image(probability_chart, width=17.0 * cm, height=7.5 * cm))
    story.extend(
        [
            Paragraph("Como interpretar os indicadores", styles["SellasH1"]),
            _pdf_table(
                ["Indicador", "Como interpretar"],
                _indicator_rows(payload),
                [4.2 * cm, 11.6 * cm],
                styles,
            ),
            Paragraph("Como os cálculos foram feitos", styles["SellasH1"]),
        ]
    )
    for index, item in enumerate(_methodology_lines(payload), start=1):
        story.append(Paragraph(f"{index}. {_escape(item)}", styles["SellasBody"]))
    ai_text = _clean_ai_text(payload.get("ai_result"))
    if ai_text:
        story.append(
            KeepTogether(
                [
                Paragraph("Leitura funcional assistida já gerada", styles["SellasH1"]),
                Paragraph(_escape(ai_text), styles["SellasBody"]),
                ]
            )
        )
    if include_limitations:
        story.extend(
            [
                Paragraph("Limitações da análise", styles["SellasH1"]),
                Paragraph(
                    "Os resultados deste relatório são descritivos e dependem da qualidade dos registros inseridos. "
                    "A associação entre antecedentes, comportamentos e consequências não demonstra causalidade nem confirma função comportamental. "
                    "A identificação de cadeias representa uma sequência temporal observada e precisa ser analisada em conjunto com avaliação funcional, "
                    "observação clínica, definição operacional, integridade dos registros e julgamento profissional. "
                    "As associações e cadeias não substituem avaliação funcional conduzida por profissional qualificado.",
                    styles["SellasBody"],
                ),
                Paragraph(
                    "Este relatório não deve ser utilizado isoladamente para prescrever intervenções, restringir oportunidades ou atribuir responsabilidade ao paciente, familiares ou profissionais.",
                    styles["SellasWarning"],
                ),
                Paragraph(
                    "As taxas de exposição não incluem automaticamente tempo em demanda ou em transição. A generalização para outros pacientes, observadores, ambientes e períodos permanece desconhecida; o PSI treino-teste não substitui monitoramento prospectivo de drift.",
                    styles["SellasBody"],
                ),
            ]
        )
    story.extend(
        [
            PageBreak(),
            Paragraph("Apêndice: lista completa das cadeias", styles["SellasH1"]),
            Paragraph(
                "Cadeias abaixo do suporte mínimo permanecem disponíveis para auditoria, mas não são destacadas no corpo principal.",
                styles["SellasBody"],
            ),
            _pdf_table(
                ["Cadeia", "k", "n", "Freq.", "IC inf.", "IC sup.", "NC"],
                _all_chain_rows(payload),
                [7.0 * cm, 1.0 * cm, 1.0 * cm, 1.7 * cm, 1.7 * cm, 1.7 * cm, 1.1 * cm],
                styles,
            ),
            Spacer(1, 8),
            Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}.", styles["SellasBody"]),
        ]
    )
    document.build(story, onFirstPage=_draw_pdf_footer, onLaterPages=_draw_pdf_footer)
    return output.getvalue()


def build_abc_report_pdf(payload: dict[str, Any]) -> bytes:
    summary = _normalized_report_summary(payload)
    narrative = build_behavior_analyst_narrative(payload)
    include_charts = bool(payload.get("include_charts", True))
    include_limitations = bool(payload.get("include_limitations", True))
    distribution_chart = _build_summary_distribution_chart(summary) if include_charts else None
    chain_chart = _build_chain_frequency_chart(payload, summary) if include_charts else None
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.35 * cm,
        leftMargin=1.35 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.25 * cm,
        title="Relatório descritivo de análise ABC",
        author="Sellas Project",
    )
    styles = _pdf_styles()
    metadata = summary.get("report_metadata") or {}
    patient = summary.get("patient") or {}
    story = [
        Paragraph("Relatório descritivo de análise ABC", styles["SellasTitle"]),
        Paragraph("Formulação clínica orientada ao analista do comportamento", styles["SellasFlow"]),
        Paragraph(
            f"<b>Paciente:</b> {_escape(patient.get('display_name') or payload.get('patient', '-'))}<br/>"
            f"<b>Ambiente:</b> {_escape(payload.get('environment') or 'Todos os ambientes')}<br/>"
            f"<b>Período:</b> {_escape(_summary_period_label(summary, payload))}<br/>"
            f"<b>Responsável pela geração:</b> {_escape(metadata.get('generated_by') or 'Usuário local')}<br/>"
            f"<b>Gerado em:</b> {_escape(_generated_at_label(metadata.get('generated_at')))}",
            styles["SellasBody"],
        ),
        Paragraph("1. Identificação e contexto disponível", styles["SellasH1"]),
        _pdf_table(["Campo", "Valor"], _clinical_identification_rows(payload), [4.5 * cm, 11.3 * cm], styles),
        Paragraph(_escape(narrative["context"]), styles["SellasBody"]),
        Paragraph("2. Síntese clínica do período", styles["SellasH1"]),
        Paragraph("Contexto relatado", styles["SellasH2"]),
        Paragraph(
            "Não há, nesta fonte, relatos subjetivos do paciente, da família ou da equipe escolar. A ausência é "
            "explicitada para evitar que inferências sejam apresentadas como fatos clínicos.", styles["SellasBody"]),
        Paragraph("Observações diretas", styles["SellasH2"]),
        Paragraph(_escape(narrative["objective"]), styles["SellasBody"]),
        Paragraph("Formulação analítico-comportamental", styles["SellasH2"]),
    ]
    for item in narrative["assessment"]:
        story.append(Paragraph(_escape(item), styles["SellasBody"]))
    story.append(Paragraph("Próximos passos de análise e mensuração", styles["SellasH2"]))
    for item in narrative["measurement_plan"]:
        story.append(Paragraph(f"• {_escape(item)}", styles["SellasBody"]))

    story.extend([
        Paragraph("3. Distribuição da observação", styles["SellasH1"]),
        Paragraph(
            "A distribuição por ambiente delimita onde a amostra foi construída. Diferenças de taxa devem ser interpretadas "
            "junto da duração, das atividades e das oportunidades observadas em cada contexto.", styles["SellasBody"]),
        _pdf_table(["Ambiente", "Horas", "Episódios", "Episódios/h"], _exposure_rows(summary),
                   [7.4 * cm, 2.6 * cm, 2.8 * cm, 3.0 * cm], styles),
    ])
    if include_charts and distribution_chart is not None:
        story.extend([
            Spacer(1, 8), Image(distribution_chart, width=17.2 * cm, height=11.0 * cm),
            Paragraph("Contagens observadas por categoria; não representam importância clínica isoladamente.", styles["SellasBody"]),
        ])

    story.extend([
        PageBreak(),
        Paragraph("4. Padrões ABC clinicamente prioritários", styles["SellasH1"]),
        Paragraph(
            "A repetição organiza a prioridade de observação. A sequência não demonstra que o antecedente causou a resposta "
            "nem que o evento posterior a reforçou.", styles["SellasBody"]),
        _pdf_table(
            ["Contexto antecedente", "Resposta observada", "Evento posterior", "Reg.", "Leitura"],
            _clinical_chain_rows(payload),
            [3.4 * cm, 3.3 * cm, 3.3 * cm, 1.2 * cm, 4.6 * cm], styles),
    ])
    if include_charts and chain_chart is not None:
        story.extend([
            Spacer(1, 8), Image(chain_chart, width=17.0 * cm, height=7.0 * cm),
            Paragraph("Frequência das cadeias priorizadas; não é uma estimativa de risco.", styles["SellasBody"]),
        ])

    story.extend([
        PageBreak(),
        Paragraph("5. Hipóteses funcionais concorrentes", styles["SellasH1"]),
        Paragraph(
            "Os rótulos funcionais registrados orientam perguntas de avaliação. Eles não confirmam a função operante de uma resposta.",
            styles["SellasBody"]),
    ])
    for item in narrative["hypotheses"]:
        story.append(Paragraph(f"• {_escape(item)}", styles["SellasBody"]))
    story.extend([
        Paragraph("O que precisa ser contrastado", styles["SellasH2"]),
        Paragraph(
            "Para cada hipótese, comparar condições relevantes, oportunidades com e sem ocorrência, consequências produzidas, "
            "latência, duração, intensidade observável, recuperação e retorno à participação.", styles["SellasBody"]),
        Paragraph("6. Organização temporal dos episódios", styles["SellasH1"]),
        Paragraph(
            "As transições abaixo descrevem episódios próximos no tempo. Elas ajudam a formular perguntas sobre recuperação e "
            "retorno à atividade, mas não demonstram uma contingência de reforçamento.", styles["SellasBody"]),
        Paragraph(
            "Nos registros legados, 'Ignorado' é preservado como valor oficial e significa que não houve mudança ambiental "
            "específica registrada; não significa omissão de cuidado.", styles["SellasBody"]),
        _pdf_table(
            ["Transição", "n", "Sess.", "Med.", "Mín.", "Máx.", "Estabilidade", "Status"],
            _clinical_temporal_rows(summary),
            [5.8 * cm, 0.8 * cm, 1.1 * cm, 1.2 * cm, 1.2 * cm, 1.2 * cm, 2.3 * cm, 2.2 * cm], styles),
    ])
    if include_limitations:
        story.extend([
            Paragraph("7. Limites, segurança e uso responsável", styles["SellasH1"]),
            Paragraph(_escape(narrative["clinical_boundary"]), styles["SellasWarning"]),
            Paragraph(
                "A linguagem descreve respostas observáveis e barreiras de participação sem atribuir culpa, intenção ou traço pessoal. "
                "Os achados dependem da completude dos registros e da representatividade do período. Qualquer procedimento deve "
                "ser definido, consentido e acompanhado por profissional qualificado, com critérios de segurança e benefício.",
                styles["SellasBody"]),
        ])

    story.extend([
        PageBreak(),
        Paragraph("Apêndice A — Auditoria dos dados e método", styles["SellasH1"]),
        Paragraph(
            "Esta seção preserva a rastreabilidade técnica. Indicadores de banco de dados, estimativas e fórmulas não devem "
            "ser lidos isoladamente como conclusões clínicas.", styles["SellasBody"]),
        _pdf_table(["Indicador", "Resultado"], _quality_rows(summary), [9.3 * cm, 6.5 * cm], styles),
        Paragraph("Indicadores descritivos", styles["SellasH2"]),
        _pdf_table(["Indicador", "Resultado"], _descriptive_metric_rows(payload), [9.3 * cm, 6.5 * cm], styles),
        Paragraph("Método de cálculo", styles["SellasH2"]),
    ])
    for index, item in enumerate(_methodology_lines(payload), start=1):
        story.append(Paragraph(f"{index}. {_escape(item)}", styles["SellasBody"]))
    story.extend([
        Paragraph("Apêndice B — Lista completa das cadeias", styles["SellasH1"]),
        Paragraph(
            "Cadeias raras são mantidas para auditoria, sem receber o mesmo destaque clínico das sequências repetidas. "
            "Os rótulos legados permanecem reproduzidos; 'Ignorado' designa ausência de mudança ambiental específica registrada.",
            styles["SellasBody"]),
        _pdf_table(
            ["Cadeia", "k", "n", "Freq.", "IC inf.", "IC sup.", "NC"],
            _all_chain_rows(payload),
            [7.0 * cm, 1.0 * cm, 1.0 * cm, 1.7 * cm, 1.7 * cm, 1.7 * cm, 1.1 * cm], styles),
        Spacer(1, 8),
        Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}.", styles["SellasBody"]),
    ])
    document.build(story, onFirstPage=_draw_pdf_footer, onLaterPages=_draw_pdf_footer)
    return output.getvalue()


def _build_overview_chart(payload: dict[str, Any]) -> io.BytesIO:
    analysis = payload.get("analysis") or {}
    series = pd.DataFrame(analysis.get("serie_temporal") or [])
    summary = _normalized_report_summary(payload)
    exposure = summary.get("exposure_summary") or (summary.get("observation_summary") or {}).get("exposure") or {}
    rate_rows = [
        {"data": day.get("date"), "comportamento": behavior, "quantidade": rate}
        for day in exposure.get("by_date") or []
        for behavior, rate in (day.get("behavior_rates_per_hour") or {}).items()
        if rate is not None
    ]
    rate_mode = bool(rate_rows)
    if rate_mode:
        series = pd.DataFrame(rate_rows)
    risk = pd.DataFrame((payload.get("analysis_all") or {}).get("mapa_risco") or [])
    risk_environment = str(payload.get("risk_environment") or "")
    if not risk.empty and risk_environment:
        risk = risk[risk["ambiente"] == risk_environment].copy()

    fig, axes = plt.subplots(1, 2, figsize=(11.4, 5.1), gridspec_kw={"width_ratios": [1.18, 1]})
    fig.patch.set_facecolor(SURFACE)
    for axis in axes:
        axis.set_facecolor(SURFACE)
        axis.grid(axis="y", color="#ded0b8", alpha=0.60, linewidth=0.8)
        axis.spines[["top", "right"]].set_visible(False)
        axis.spines[["left", "bottom"]].set_color("#b9aa90")
        axis.tick_params(labelsize=8, colors=TEXT)

    if series.empty:
        axes[0].text(0.5, 0.5, "Sem série temporal disponível", ha="center", va="center")
    else:
        pivot = series.pivot_table(index="data", columns="comportamento", values="quantidade", aggfunc="sum", fill_value=0)
        pivot = pivot.sort_index()
        if len(pivot.columns) > 6:
            keep = pivot.sum().sort_values(ascending=False).head(6).index
            pivot = pivot[keep]
        bottom = pd.Series(0.0, index=pivot.index)
        palette = [RED, GREEN, BLUE, GOLD, "#8a7a9d", "#6f9b96"]
        x = range(len(pivot.index))
        for index, column in enumerate(pivot.columns):
            values = pivot[column].astype(float)
            axes[0].bar(x, values, bottom=bottom, label=_wrap_chart_label(column, 22), color=palette[index % len(palette)])
            bottom += values
        labels = [pd.to_datetime(value, errors="coerce").strftime("%d/%m") for value in pivot.index]
        axes[0].set_xticks(list(x), labels, rotation=45, ha="right")
        axes[0].legend(frameon=False, fontsize=7, loc="upper left")
    axes[0].set_title("Taxa por data" if rate_mode else "Contagem por data (sem exposição)", loc="left", fontsize=12, fontweight="bold", color=TEXT)
    axes[0].set_ylabel("Episódios por hora observada" if rate_mode else "Registros", color=TEXT)

    axes[1].axvspan(0, 50, color="#dce8d8", alpha=0.30)
    axes[1].axvspan(50, 100, color="#ede1b8", alpha=0.25)
    axes[1].axvline(50, color="#8f8171", linestyle="--", linewidth=0.9)
    if risk.empty:
        axes[1].text(0.5, 0.5, "Sem cadeias no local selecionado", ha="center", va="center", transform=axes[1].transAxes)
    else:
        marker_map = {"C1": "o", "C2": "^", "NC": "x"}
        color_map = {"C1": GREEN, "C2": RED, "NC": "#6f6b66"}
        severity_values = [
            _finite(row.get("peso_medio_gravidade", row.get("indice_perigo")))
            for _, row in risk.head(12).iterrows()
        ]
        annotated_behaviors: set[str] = set()
        label_offsets = [(0, 8), (0, 24), (0, 40), (0, 56), (0, 72), (0, 88)]
        for _, row in risk.head(12).iterrows():
            classification = str(row.get("classificacao_predominante") or "NC")
            x_value = _finite(row.get("probabilidade_ocorrencia")) * 100
            severity_value = row.get("peso_medio_gravidade", row.get("indice_perigo"))
            y_value = _finite(severity_value) if severity_value is not None else 0.0
            size = 45 + min(160, _finite(row.get("suporte")) * 8)
            marker = marker_map.get(classification, "x")
            scatter_options = {
                "s": size,
                "marker": marker,
                "color": color_map.get(classification, "#6f6b66"),
                "linewidths": 1.2,
            }
            if marker != "x":
                scatter_options["edgecolors"] = "white"
            axes[1].scatter(x_value, y_value, **scatter_options)
            behavior_label = str(row.get("comportamento") or "Não identificado")
            if behavior_label not in annotated_behaviors and len(annotated_behaviors) < len(label_offsets):
                offset_x, offset_y = label_offsets[len(annotated_behaviors)]
                axes[1].annotate(
                    _wrap_chart_label(behavior_label, 14),
                    (x_value, y_value),
                    xytext=(offset_x, offset_y),
                    textcoords="offset points",
                    ha="center",
                    fontsize=7,
                    color=TEXT,
                )
                annotated_behaviors.add(behavior_label)
        if len({round(value, 6) for value in severity_values}) <= 1:
            axes[1].text(
                0.02,
                0.98,
                "Gravidade sem variação neste recorte;\no eixo vertical não diferencia cadeias.",
                transform=axes[1].transAxes,
                va="top",
                fontsize=7,
                color="#8a4a42",
            )
    axes[1].set_xlim(-2, 102)
    axes[1].set_ylim(-0.02, 1.02)
    axes[1].set_xlabel("Frequência observada da cadeia no local (%)", color=TEXT)
    axes[1].set_ylabel("Peso médio configurado de gravidade", color=TEXT)
    axes[1].set_title(f"Frequência × gravidade configurada: {_wrap_chart_label(risk_environment or 'local', 24)}", loc="left", fontsize=12, fontweight="bold", color=TEXT)
    fig.tight_layout(pad=1.2)
    return _figure_bytes(fig)


def _build_probability_chart(payload: dict[str, Any]) -> io.BytesIO:
    prediction = payload.get("prediction") or {}
    chain = payload.get("selected_chain") or {}
    mode = str(prediction.get("analysis_mode") or "descriptive")
    labels = ["P(A-B-C)", "P(B|A)", "P(C|A,B)"]
    values = [
        _finite(chain.get("probabilidade_conjunta")) * 100,
        _finite(chain.get("probabilidade_comportamento_dado_antecedente")) * 100,
        _finite(chain.get("probabilidade_consequencia_dada_cadeia_ab")) * 100,
    ]
    if mode == "predictive" and prediction.get("predicted_probability") is not None:
        labels.append("Probabilidade prevista")
        values.append(_finite(prediction.get("predicted_probability")) * 100)
    elif prediction.get("estimativa_descritiva", prediction.get("probabilidade_prevista")) is not None:
        labels.append("Frequência do alvo")
        values.append(_finite(prediction.get("estimativa_descritiva", prediction.get("probabilidade_prevista"))) * 100)
    fig, axis = plt.subplots(figsize=(10.8, 4.4))
    fig.patch.set_facecolor(SURFACE)
    axis.set_facecolor(SURFACE)
    bars = axis.bar(labels, values, color=[GOLD, GREEN, RED, BLUE])
    axis.set_ylim(0, 108)
    axis.set_ylabel("Proporção observada (%)" if mode != "predictive" else "Probabilidade (%)", color=TEXT)
    chart_title = "Proporções observadas da cadeia e do alvo" if mode != "predictive" else "Probabilidades observadas e previsão validada"
    axis.set_title(chart_title, loc="left", fontsize=13, fontweight="bold", color=TEXT)
    axis.grid(axis="y", color="#ded0b8", alpha=0.65, linewidth=0.8)
    axis.spines[["top", "right"]].set_visible(False)
    axis.tick_params(axis="x", labelrotation=18, labelsize=8)
    for bar, value in zip(bars, values, strict=False):
        axis.text(bar.get_x() + bar.get_width() / 2, min(104, value + 2), f"{value:.1f}%", ha="center", fontsize=8, color=TEXT)
    fig.tight_layout(pad=1.2)
    return _figure_bytes(fig)


def _build_summary_distribution_chart(summary: dict[str, Any]) -> io.BytesIO:
    fig, axes = plt.subplots(1, 3, figsize=(12.2, 5.2))
    fig.patch.set_facecolor("white")
    panels = [
        ("Antecedentes", summary.get("top_antecedents") or [], GREEN),
        ("Comportamentos", summary.get("top_behaviors") or [], RED),
        ("Consequências", summary.get("top_consequences") or [], BLUE),
    ]
    for axis, (title, rows, color) in zip(axes, panels, strict=False):
        axis.set_facecolor("white")
        visible = list(reversed(rows[:7]))
        if not visible:
            axis.text(0.5, 0.5, "Sem dados", ha="center", va="center", transform=axis.transAxes)
        else:
            labels = [_wrap_chart_label(item.get("nome"), 24) for item in visible]
            values = [int(item.get("quantidade") or 0) for item in visible]
            bars = axis.barh(labels, values, color=color, edgecolor="#ffffff")
            for bar, value in zip(bars, values, strict=False):
                axis.text(value + max(values) * 0.02, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=8)
            axis.set_xlim(0, max(values) * 1.18 if max(values) else 1)
        axis.set_title(title, loc="left", fontsize=11, fontweight="bold", color=TEXT)
        axis.grid(axis="x", color="#ded0b8", alpha=0.55)
        axis.spines[["top", "right", "left"]].set_visible(False)
        axis.tick_params(labelsize=8, colors=TEXT)
    fig.tight_layout(pad=1.4)
    return _figure_bytes(fig)


def _build_association_heatmap(summary: dict[str, Any]) -> io.BytesIO:
    fig, axes = plt.subplots(1, 2, figsize=(12.0, 5.6))
    fig.patch.set_facecolor("white")

    def render(axis, rows: list[dict[str, Any]], left_key: str, right_key: str, title: str) -> None:
        visible = rows[:12]
        if not visible:
            axis.text(0.5, 0.5, "Sem dados suficientes", ha="center", va="center", transform=axis.transAxes)
            axis.set_title(title, loc="left", fontsize=11, fontweight="bold", color=TEXT)
            axis.set_axis_off()
            return
        left_values = list(dict.fromkeys(str(item.get(left_key) or "-") for item in visible))[:6]
        right_values = list(dict.fromkeys(str(item.get(right_key) or "-") for item in visible))[:6]
        matrix = []
        for left in left_values:
            row_values = []
            for right in right_values:
                item = next(
                    (
                        entry
                        for entry in visible
                        if str(entry.get(left_key) or "-") == left and str(entry.get(right_key) or "-") == right
                    ),
                    None,
                )
                row_values.append(_finite(item.get("probabilidade_condicional")) * 100 if item else 0.0)
            matrix.append(row_values)
        image = axis.imshow(matrix, cmap="YlGn", vmin=0, vmax=100, aspect="auto")
        axis.set_xticks(range(len(right_values)), [_wrap_chart_label(value, 16) for value in right_values], rotation=35, ha="right")
        axis.set_yticks(range(len(left_values)), [_wrap_chart_label(value, 20) for value in left_values])
        for y, row_values in enumerate(matrix):
            for x, value in enumerate(row_values):
                axis.text(x, y, f"{value:.0f}%", ha="center", va="center", fontsize=7, color="#1f1f1f")
        axis.set_title(title, loc="left", fontsize=11, fontweight="bold", color=TEXT)
        axis.tick_params(labelsize=7)
        fig.colorbar(image, ax=axis, fraction=0.046, pad=0.04, label="Proporção condicional observada (%)")

    render(
        axes[0],
        summary.get("antecedent_behavior_associations") or [],
        "antecedente",
        "comportamento",
        "P(comportamento | antecedente)",
    )
    render(
        axes[1],
        summary.get("behavior_consequence_associations") or [],
        "comportamento",
        "consequencia",
        "P(consequência | comportamento)",
    )
    fig.tight_layout(pad=1.3)
    return _figure_bytes(fig)


def _build_chain_frequency_chart(payload: dict[str, Any], summary: dict[str, Any]) -> io.BytesIO:
    chains = sorted(
        [item for item in ((payload.get("analysis") or {}).get("cadeias_completas") or []) if int(item.get("suporte") or 0) >= 8],
        key=lambda item: int(item.get("suporte") or 0),
        reverse=True,
    )[:10]
    fig, axis = plt.subplots(figsize=(11.2, 4.5))
    fig.patch.set_facecolor("white")
    axis.set_facecolor("white")
    if not chains:
        axis.text(0.5, 0.5, "Sem cadeias A-B-C no recorte", ha="center", va="center", transform=axis.transAxes)
        axis.set_axis_off()
    else:
        visible = list(reversed(chains))
        labels = [_wrap_chart_label(item.get("cadeia"), 46) for item in visible]
        values = [int(item.get("suporte") or 0) for item in visible]
        bars = axis.barh(labels, values, color=GOLD, edgecolor="#9d8955")
        for bar, value in zip(bars, values, strict=False):
            axis.text(value + max(values) * 0.015, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=8)
        axis.set_xlim(0, max(values) * 1.12 if max(values) else 1)
        axis.set_xlabel("Repetições observadas", color=TEXT)
        axis.grid(axis="x", color="#ded0b8", alpha=0.55)
        axis.spines[["top", "right", "left"]].set_visible(False)
        axis.tick_params(labelsize=8, colors=TEXT)
    axis.set_title("Principais cadeias A-B-C", loc="left", fontsize=12, fontweight="bold", color=TEXT)
    fig.tight_layout(pad=1.2)
    return _figure_bytes(fig)


def _normalized_report_summary(payload: dict[str, Any]) -> dict[str, Any]:
    summary = payload.get("report_summary")
    if isinstance(summary, dict) and summary:
        return summary
    analysis = payload.get("analysis") or {}
    total = int(analysis.get("total_registros") or 0)
    top_chain = (analysis.get("cadeias_completas") or [{}])[0]
    return {
        "report_metadata": {
            "generated_by": "Usuário local",
            "instrument_version": "1",
            "logic_version": "abc-print-summary-v1",
        },
        "patient": {"display_name": payload.get("patient") or "-", "anonymized": False},
        "observation_summary": {
            "sessions": total,
            "total_intervals": total,
            "observed_intervals": total,
            "not_observed_intervals": 0,
            "invalid_intervals": 0,
            "coverage": 1.0 if total else 0.0,
            "observed_hours": total * 5 / 60,
            "occurrence_intervals": total,
            "occurrence_percentage": 1.0 if total else 0.0,
            "missing_percentage": 0.0,
        },
        "most_frequent": {
            "antecedent": {"nome": top_chain.get("antecedente") or "Não disponível", "quantidade": top_chain.get("suporte") or 0},
            "behavior": {"nome": top_chain.get("comportamento") or "Não disponível", "quantidade": top_chain.get("suporte") or 0},
            "consequence": {"nome": top_chain.get("consequencia") or "Não disponível", "quantidade": top_chain.get("suporte") or 0},
            "chain": top_chain,
        },
        "top_antecedents": [],
        "top_behaviors": [],
        "top_consequences": [],
        "antecedent_behavior_associations": analysis.get("antecedente_comportamento") or [],
        "behavior_consequence_associations": analysis.get("comportamento_consequencia") or [],
        "behavior_chains": [],
        "data_quality": {"status": "inicial" if total else "insuficiente", "coverage": 1.0 if total else 0.0, "warnings": []},
        "descriptive_summary": build_abc_plain_language_explanation(payload)["headline"],
        "clinical_disclaimer": DEFAULT_CLINICAL_DISCLAIMER,
    }


def _clinical_summary_rows(summary: dict[str, Any], payload: dict[str, Any]) -> list[list[str]]:
    observation = summary.get("observation_summary") or {}
    frequent = summary.get("most_frequent") or {}
    chain = frequent.get("chain") or {}
    temporal = summary.get("behavior_chains") or []
    mean_time = [item.get("tempo_medio_segundos") for item in temporal if item.get("tempo_medio_segundos") is not None]
    return [
        ["Sessões", _display_value(observation.get("sessions"))],
        ["Registros ABC completos", _display_value(observation.get("total_records"))],
        ["Ocorrências de comportamento", _display_value(observation.get("behavior_occurrences"))],
        ["Horas observadas", f"{_finite(observation.get('observed_hours')):.2f} h"],
        ["Intervalos totais / observados", f"{int(observation.get('total_intervals') or 0)} / {int(observation.get('observed_intervals') or 0)}"],
        ["Cobertura observacional", _pct(_finite(observation.get("coverage")))],
        ["Intervalos com ocorrência", f"{int(observation.get('occurrence_intervals') or 0)} ({_pct(_finite(observation.get('occurrence_percentage')))})"],
        ["Antecedente mais frequente", str((frequent.get("antecedent") or {}).get("nome") or "Não disponível")],
        ["Comportamento mais frequente", str((frequent.get("behavior") or {}).get("nome") or "Não disponível")],
        ["Consequência mais frequente", str((frequent.get("consequence") or {}).get("nome") or "Não disponível")],
        ["Cadeia A-B-C mais frequente", str(chain.get("cadeia") or "Não disponível")],
        ["Tempo médio entre elos temporais", f"{sum(mean_time) / len(mean_time):.1f} s" if mean_time else "Não disponível"],
    ]


def _clinical_identification_rows(payload: dict[str, Any]) -> list[list[str]]:
    summary = _normalized_report_summary(payload)
    observation = summary.get("observation_summary") or {}
    return [
        ["Paciente", str((summary.get("patient") or {}).get("display_name") or payload.get("patient") or "-")],
        ["Período analisado", _summary_period_label(summary, payload)],
        ["Ambientes incluídos", str(payload.get("environment") or "Todos os ambientes registrados")],
        ["Fonte", "Registros diretos no instrumento ABC fechado"],
        ["Volume da amostra", f"{int(observation.get('total_records') or (payload.get('analysis') or {}).get('total_registros') or 0)} registros; {_decimal_pt(observation.get('observed_hours'), 1)} h observadas"],
        ["Dados biopsicossociais", "Não informados nesta fonte; integrar antes de decisões clínicas"],
    ]


def _clinical_chain_rows(payload: dict[str, Any]) -> list[list[str]]:
    chains = sorted(
        (payload.get("analysis") or {}).get("cadeias_completas") or [],
        key=lambda item: int(item.get("suporte") or 0),
        reverse=True,
    )[:5]
    rows = []
    for item in chains:
        support = int(item.get("suporte") or 0)
        rows.append([
            str(item.get("antecedente") or "Não informado"),
            str(item.get("comportamento") or "Não informado"),
            str(item.get("consequencia") or "Não informado"),
            str(support),
            "Priorizar observação dirigida" if support >= 8 else "Manter como padrão emergente",
        ])
    return rows or [["Sem padrão priorizável", "-", "-", "0", "Ampliar observação"]]


def _clinical_temporal_rows(summary: dict[str, Any]) -> list[list[str]]:
    rows = []
    for item in (summary.get("behavior_chains") or [])[:8]:
        rows.append([
            _humanize_temporal_chain(item.get("cadeia")),
            str(int(item.get("repeticoes") or 0)),
            _display_value(item.get("sessoes")),
            _seconds_label(item.get("tempo_mediano_segundos", item.get("tempo_medio_segundos"))),
            _seconds_label(item.get("tempo_minimo_segundos")),
            _seconds_label(item.get("tempo_maximo_segundos")),
            str(item.get("estabilidade_temporal") or "inicial"),
            str(item.get("status") or "Candidata"),
        ])
    return rows or [["Nenhuma transição temporal válida", "0", "-", "-", "-", "-", "insuficiente", "-"]]


def _humanize_temporal_chain(value: Any) -> str:
    text = str(value or "-")
    replacements = {
        "AGRESSAO_FISICA": "Agressão física",
        "AUTOLESAO": "Autolesão",
        "SE_JOGAR_NO_CHAO": "Se jogar no chão",
        "BATER_OS_PES_NA_PAREDE": "Bater os pés na parede",
        "COMPORTAMENTO_BLOQUEADO": "Comportamento bloqueado",
        "MANEJO_FISICO": "Manejo físico",
        "DO_NADA": "Contexto não operacionalizado ('do nada')",
        "IGNORADO": "Sem mudança ambiental registrada",
        "DEMANDA": "Demanda apresentada",
    }
    for prefix in ("COM_", "CON_", "ANT_"):
        text = text.replace(prefix, "")
    for code, label in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(code, label)
    text = text.replace(" => ", " | próximo episódio: ").replace(" -> ", " → ")
    return text


def _abc_component_rows() -> list[list[str]]:
    return [
        ["Antecedente", "Evento, situação, atividade ou condição observada antes do comportamento, como demanda, espera, transição, retirada de item ou pedido negado."],
        ["Comportamento", "Resposta observável e operacionalmente definida, como agressão física, autolesão, fuga, protesto vocal, jogar objetos ou chorar."],
        ["Consequência", "Evento ou mudança ambiental observada depois do comportamento, como retirada da demanda, acesso a item, atenção, redirecionamento, continuidade ou pausa."],
    ]


def _observation_state_rows() -> list[list[str]]:
    return [
        ["Ocorreu", "O evento foi observado no intervalo."],
        ["Não ocorreu", "O intervalo foi observado e o evento não ocorreu."],
        ["Não informado", "Não foi possível determinar o valor do evento."],
        ["Não observado", "O intervalo não teve observação válida e fica fora dos denominadores."],
        ["Inválido", "Registro excluído da análise por inconsistência."],
    ]


def _quality_rows(summary: dict[str, Any]) -> list[list[str]]:
    quality = summary.get("data_quality") or {}
    observation = summary.get("observation_summary") or {}
    return [
        ["Classificação geral", str(quality.get("status_scope") or quality.get("status") or "insuficiente").capitalize()],
        ["Prontidão preditiva", str(quality.get("predictive_readiness") or "Não avaliada")],
        ["Cobertura", _pct(_finite(quality.get("coverage")))],
        ["Horas observadas", _decimal_pt(observation.get("observed_hours"), 1)],
        ["Episódios por hora observada", _decimal_pt(observation.get("occurrences_per_observed_hour"), 1)],
        ["Oportunidades observadas sem episódio", str(int(observation.get("opportunities_without_occurrence") or 0))],
        ["Intervalos não observados", str(int(observation.get("not_observed_intervals") or 0))],
        ["Intervalos inválidos", str(int(observation.get("invalid_intervals") or 0))],
        ["Percentual ausente", _pct(_finite(observation.get("missing_percentage")))],
        ["Valores de evento não informados", str(int(quality.get("missing_event_values") or 0))],
        ["Eventos sem revisão", str(int(quality.get("unreviewed_events") or 0))],
        ["Registros retroativos", str(int(quality.get("retroactive_records") or 0))],
        ["Duplicidades detectadas", str(int(quality.get("duplicates") or 0))],
        ["Timestamps ausentes / inválidos", f"{int(quality.get('incomplete_times') or 0)} / {int(quality.get('invalid_timestamps') or 0)}"],
        ["Sobreposições temporais", str(int(quality.get("overlapping_intervals") or 0)) if quality.get("overlap_evaluable") else "Não avaliável"],
        ["Categorias fora da taxonomia", str(int(quality.get("outside_taxonomy") or 0)) if quality.get("taxonomy_evaluable") else "Não avaliável"],
        ["Observadores identificados", str(int(quality.get("observer_count") or 0)) if quality.get("observer_identification_available") else "Não avaliável"],
        ["Registros sem identificador", str(int(quality.get("missing_record_identifiers") or 0))],
        ["Cadeias censuradas", str(int(quality.get("censored_chains") or 0))],
    ]


def _exposure_rows(summary: dict[str, Any]) -> list[list[str]]:
    exposure = summary.get("exposure_summary") or (summary.get("observation_summary") or {}).get("exposure") or {}
    rows = []
    for environment, values in sorted((exposure.get("by_environment") or {}).items()):
        rows.append(
            [
                str(environment),
                _decimal_pt(values.get("observed_hours"), 1),
                str(int(values.get("occurrence_intervals") or 0)),
                _decimal_pt(values.get("occurrences_per_hour"), 1),
            ]
        )
    return rows or [["Exposição não disponível", "-", "-", "-"]]


def _association_table_rows(summary: dict[str, Any], key: str, left_key: str, right_key: str) -> list[list[str]]:
    rows = []
    for item in (summary.get(key) or [])[:8]:
        lift = item.get("lift")
        rows.append(
            [
                _short(item.get(left_key), 30),
                _short(item.get(right_key), 30),
                str(int(item.get("exposicoes") or item.get("total_exposto") or 0)),
                str(int(item.get("ocorrencias_conjuntas") or item.get("suporte") or 0)),
                _pct(_finite(item.get("probabilidade_condicional"))),
                _pct(_finite(item.get("probabilidade_baseline"))),
                f"{_finite(item.get('diferenca_risco')) * 100:+.1f} p.p.",
                f"{_finite(lift):.2f}" if lift is not None else "-",
                str(item.get("qualidade_estimativa") or "inicial"),
            ]
        )
    return rows or [["Sem dados", "-", "0", "0", "-", "-", "-", "-", "insuficiente"]]


def _temporal_summary_rows(summary: dict[str, Any]) -> list[list[str]]:
    rows = []
    for item in summary.get("behavior_chains") or []:
        rows.append(
            [
                str(item.get("cadeia") or "-"),
                str(int(item.get("repeticoes") or 0)),
                _display_value(item.get("sessoes")),
                _seconds_label(item.get("tempo_medio_segundos")),
                _seconds_label(item.get("tempo_minimo_segundos")),
                _seconds_label(item.get("tempo_maximo_segundos")),
                str(item.get("estabilidade_temporal") or "inicial"),
                _pct(_finite(item.get("confianca_media"))) if item.get("confianca_media") is not None else "-",
                str(item.get("status") or "Candidata"),
            ]
        )
    return rows or [["Nenhuma cadeia temporal válida no filtro", "0", "-", "-", "-", "-", "-", "-", "-"]]


def _summary_period_label(summary: dict[str, Any], payload: dict[str, Any]) -> str:
    metadata = summary.get("report_metadata") or {}
    start = metadata.get("period_start")
    end = metadata.get("period_end")
    if start and end:
        try:
            return f"{pd.to_datetime(start).strftime('%d/%m/%Y')} a {pd.to_datetime(end).strftime('%d/%m/%Y')}"
        except Exception:
            pass
    return _period_label(payload)


def _generated_at_label(value: Any) -> str:
    if not value:
        return "Não disponível"
    try:
        timestamp = pd.to_datetime(value)
        return timestamp.strftime("%d/%m/%Y %H:%M")
    except Exception:
        return str(value)


def _display_value(value: Any) -> str:
    return "Não disponível" if value is None or value == "" else str(value)


def _seconds_label(value: Any) -> str:
    return "-" if value is None else f"{_finite(value):.1f} s"


def _identification_rows(payload: dict[str, Any]) -> list[list[str]]:
    prediction = payload.get("prediction") or {}
    analysis = payload.get("analysis") or {}
    summary = _normalized_report_summary(payload)
    metadata = summary.get("report_metadata") or {}
    mode = str(prediction.get("analysis_mode") or analysis.get("analysis_mode") or "descriptive")
    return [
        ["Paciente", str((summary.get("patient") or {}).get("display_name") or payload.get("patient") or "-")],
        ["Período", _summary_period_label(summary, payload)],
        ["Ambientes", str(payload.get("environment") or "Todos os ambientes")],
        ["Ambiente da visualização", str(payload.get("risk_environment") or payload.get("environment") or "Todos os ambientes")],
        ["Unidade de análise", "Episódio ABC registrado" if mode == "descriptive" else str(prediction.get("prediction_unit") or "Não informada")],
        ["Alvo e horizonte", f"{prediction.get('comportamento') or '-'} / {prediction.get('prediction_horizon') or 'não aplicável no modo descritivo'}"],
        ["Modo analítico", "Descritivo" if mode == "descriptive" else "Preditivo validado"],
        ["Versões", f"dados {metadata.get('instrument_version') or '-'}; metodologia {analysis.get('methodology_version') or metadata.get('methodology_version') or METHODOLOGY_VERSION}; pipeline {metadata.get('pipeline_version') or '-'}"],
        ["Dataset / execução", f"{metadata.get('dataset_version') or '-'} / {str(metadata.get('analysis_run_hash') or '-')[:16]}"],
        ["Taxonomia / origem", f"{', '.join(metadata.get('taxonomy_versions') or []) or 'não informada'} / {', '.join(metadata.get('source_systems') or []) or 'não informada'}"],
        ["Geração", _generated_at_label(metadata.get("generated_at"))],
    ]


def _descriptive_metric_rows(payload: dict[str, Any]) -> list[list[str]]:
    analysis = payload.get("analysis") or {}
    prediction = payload.get("prediction") or {}
    summary = _normalized_report_summary(payload)
    estimate = prediction.get("estimativa_descritiva", prediction.get("probabilidade_prevista"))
    observation = summary.get("observation_summary") or {}
    metadata = summary.get("report_metadata") or {}
    sessions = int(analysis.get("sessoes_unicas") or observation.get("sessions") or 0)
    days = int(analysis.get("dias_unicos") or metadata.get("period_days") or 0)
    rows = [
        ["Registros incluídos", str(int(analysis.get("total_registros") or 0))],
        ["Sessões / dias", f"{sessions} / {days}"],
        ["Alvo descritivo", str(prediction.get("comportamento") or "-")],
        ["Numerador k / denominador n", f"{int(prediction.get('numerador') or prediction.get('sucessos_contexto') or 0)} / {int(prediction.get('denominador') or prediction.get('amostra_contexto') or 0)}"],
        ["Frequência observada", _pct_or_na(estimate)],
        ["Método", str(prediction.get("metodo") or "proporção bruta com IC de Wilson")],
        ["Intervalo", f"{_pct_or_na(prediction.get('intervalo_wilson_inferior'))} a {_pct_or_na(prediction.get('intervalo_wilson_superior'))}"],
        ["Sem classificação C1/C2", str(int(analysis.get("nao_classificados") or 0))],
    ]
    clustered = prediction.get("intervalo_agrupado_sessao") or {}
    clustered_label = "IC agrupado por sessão no mesmo contexto"
    if not clustered:
        clustered = summary.get("top_behavior_session_cluster_interval") or {}
        clustered_label = "IC agrupado por sessão do comportamento mais frequente no recorte"
    if clustered:
        rows.append(
            [
                clustered_label,
                f"{_pct_or_na(clustered.get('lower'))} a {_pct_or_na(clustered.get('upper'))}; {int(clustered.get('cluster_count') or 0)} sessões",
            ]
        )
    return rows


def _predictive_rows(prediction: dict[str, Any]) -> list[list[str]]:
    metrics = prediction.get("metrics") or {}
    baseline = prediction.get("baseline") or {}
    periods = prediction.get("periods") or {}
    matrix = metrics.get("confusion_matrix") or {}
    operational = prediction.get("operational_evaluation") or metrics.get("operational") or {}
    drift = prediction.get("drift") or {}
    backtest = prediction.get("rolling_origin_backtest") or {}
    explanation = prediction.get("individual_explanation") or {}
    contributions = explanation.get("top_contributions") or []
    return [
        ["Alvo / horizonte", f"{prediction.get('target') or prediction.get('comportamento') or '-'} / {prediction.get('prediction_horizon') or '-'}"],
        ["Treino", f"{periods.get('training_start') or '-'} a {periods.get('training_end') or '-'}"],
        ["Teste", f"{periods.get('validation_start') or '-'} a {periods.get('validation_end') or '-'}"],
        ["Positivos / negativos no teste", f"{prediction.get('validation_positives', '-')} / {prediction.get('validation_negatives', '-')}"],
        ["Modelo", str(prediction.get("model") or "-")],
        ["Versão do modelo / features", f"{prediction.get('model_version') or '-'} / {prediction.get('feature_version') or '-'}"],
        ["Parâmetros", "; ".join(f"{key}={value}" for key, value in (prediction.get('model_parameters') or {}).items()) or "-"],
        ["Melhor baseline", str(baseline.get("name") or "-")],
        ["Baselines comparados", str(len(prediction.get("baselines") or ([baseline] if baseline else [])))],
        ["Brier / Log Loss", f"{_number_or_na(metrics.get('brier_score'))} / {_number_or_na(metrics.get('log_loss'))}"],
        ["PR-AUC / ROC-AUC", f"{_number_or_na(metrics.get('pr_auc'))} / {_number_or_na(metrics.get('roc_auc'))}"],
        ["Precisão / recall / especificidade", f"{_number_or_na(metrics.get('precision'))} / {_number_or_na(metrics.get('recall'))} / {_number_or_na(metrics.get('specificity'))}"],
        ["Matriz TN / FP / FN / TP", f"{matrix.get('tn', '-')} / {matrix.get('fp', '-')} / {matrix.get('fn', '-')} / {matrix.get('tp', '-')}"],
        ["Alertas / falsos / perdidos", f"{operational.get('alerts', '-')} / {operational.get('false_alerts', '-')} / {operational.get('missed_events', '-')} (limiar {_number_or_na(operational.get('threshold'))})"],
        ["Calibração (intercepto / inclinação / ECE)", f"{_number_or_na(metrics.get('calibration_intercept'))} / {_number_or_na(metrics.get('calibration_slope'))} / {_number_or_na(metrics.get('expected_calibration_error'))}"],
        ["Backtesting progressivo", f"{backtest.get('status') or 'não avaliado'}; {int(backtest.get('fold_count') or 0)} janela(s)"],
        ["Drift treino-teste", f"{drift.get('status') or 'não avaliado'}; PSI máx. {_number_or_na(drift.get('max_psi'))}"],
        ["Fatores da última oportunidade", "; ".join(f"{item.get('feature')}: {float(item.get('log_odds_contribution') or 0):+.3f}" for item in contributions[:5]) or "Não disponível"],
        ["Hash da execução", str(prediction.get("analysis_run_hash") or "-")[:20]],
        ["Conclusão", str(prediction.get("conclusion") or "-")],
    ]


def _top_chain_rows_v2(payload: dict[str, Any]) -> list[list[str]]:
    minimum_support = 8
    chains = [
        item for item in sorted((payload.get("analysis") or {}).get("cadeias_completas") or [], key=lambda item: int(item.get("suporte") or 0), reverse=True)
        if int(item.get("suporte") or 0) >= minimum_support
    ][:10]
    rows = []
    for item in chains:
        rows.append(
            [
                str(item.get("cadeia") or "-"),
                f"{int(item.get('numerador_conjunto') or item.get('suporte') or 0)}/{int(item.get('denominador_conjunto') or (payload.get('analysis') or {}).get('total_registros') or 0)}",
                _pct_or_na(item.get("probabilidade_conjunta")),
                str(int(item.get("sessoes_unicas") or 0)),
                str(int(item.get("dias_unicos") or 0)),
                str(item.get("evidence_status") or "evidência insuficiente"),
            ]
        )
    return rows or [[f"Nenhuma cadeia atingiu suporte mínimo {minimum_support}; consultar apêndice", "0/0", "-", "0", "0", "insuficiente"]]


def _severity_rows(payload: dict[str, Any]) -> list[list[str]]:
    analysis = payload.get("analysis") or {}
    chain = payload.get("selected_chain") or {}
    return [
        ["C1 / C2 / não classificados", f"{int(analysis.get('c1_leve') or 0)} / {int(analysis.get('c2_intenso') or 0)} / {int(analysis.get('nao_classificados') or 0)}"],
        ["Peso configurado C1 / C2", "0,20 / 1,00 (regra abc-severity-v1)"],
        ["Peso médio da cadeia", _number_or_na(chain.get("peso_medio_gravidade", chain.get("indice_perigo")))],
        ["Índice exploratório", _number_or_na(chain.get("indice_risco_exploratorio", chain.get("indice_risco")))],
        ["Interpretação", str(chain.get("mensagem_gravidade") or "Peso interno exploratório; não é probabilidade de dano.")],
    ]


def _temporal_rows_v2(payload: dict[str, Any]) -> list[list[str]]:
    summary = _normalized_report_summary(payload)
    rows = []
    for item in (summary.get("behavior_chains") or [])[:10]:
        rows.append(
            [
                _humanize_temporal_chain(item.get("cadeia")),
                str(int(item.get("repeticoes") or 0)),
                str(item.get("sessoes") or 0),
                str(item.get("dias") or 0),
                _seconds_label(item.get("tempo_mediano_segundos", item.get("tempo_medio_segundos"))),
                _seconds_label(item.get("tempo_iqr_segundos")),
                str(item.get("estabilidade_temporal") or "evidência insuficiente"),
            ]
        )
    return rows or [["Nenhuma transição temporal válida", "0", "0", "0", "-", "-", "insuficiente"]]


def _functional_hypothesis_rows(payload: dict[str, Any]) -> list[list[str]]:
    rows = []
    for item in (payload.get("analysis") or {}).get("por_funcao") or []:
        rows.append([str(item.get("funcao") or "Não identificada"), str(int(item.get("quantidade") or 0))])
    if not rows:
        chain = payload.get("selected_chain") or {}
        if chain.get("funcao_predominante"):
            rows.append([str(chain["funcao_predominante"]), str(int(chain.get("suporte") or 0))])
    return rows or [["Não identificada", "0"]]


def _all_chain_rows(payload: dict[str, Any]) -> list[list[str]]:
    total = int((payload.get("analysis") or {}).get("total_registros") or 0)
    rows = []
    for item in sorted((payload.get("analysis") or {}).get("cadeias_completas") or [], key=lambda value: int(value.get("suporte") or 0), reverse=True):
        rows.append(
            [
                str(item.get("cadeia") or "-"),
                str(int(item.get("suporte") or 0)),
                str(int(item.get("denominador_conjunto") or total)),
                _pct_or_na(item.get("probabilidade_conjunta")),
                _pct_or_na(item.get("intervalo_inferior")),
                _pct_or_na(item.get("intervalo_superior")),
                str(int(item.get("nao_classificado") or 0)),
            ]
        )
    return rows or [["Sem cadeia", "0", str(total), "-", "-", "-", "0"]]


def _metric_rows(payload: dict[str, Any]) -> list[list[str]]:
    analysis = payload.get("analysis") or {}
    prediction = payload.get("prediction") or {}
    chain = payload.get("selected_chain") or {}
    temporal = payload.get("temporal_data") or {}
    candidates = [item for item in temporal.get("candidates") or [] if item.get("rejection_reason") != "superseded_detection"]
    accepted = sum(item.get("validation_status") == "accepted" for item in candidates)
    stable = sum(not bool(item.get("insufficient_sample")) for item in temporal.get("stats") or [])
    chain_scope = str(payload.get("chain_scope") or "selected")
    chain_label = "Cadeia de referência" if chain_scope == "all" else "Cadeia selecionada"
    rows = [
        ["Registros no recorte", str(int(analysis.get("total_registros") or 0))],
        ["C1 - leve / C2 - intenso", f"{int(analysis.get('c1_leve') or 0)} / {int(analysis.get('c2_intenso') or 0)}"],
        ["Comportamento descrito", str(prediction.get("comportamento") or "-")],
        ["Frequência observada", _pct_or_na(prediction.get("estimativa_descritiva", prediction.get("probabilidade_prevista")))],
        ["IC de Wilson da proporção", f"{_pct_or_na(prediction.get('intervalo_wilson_inferior'))} a {_pct_or_na(prediction.get('intervalo_wilson_superior'))}"],
        ["Qualidade da evidência", str(prediction.get("qualidade_evidencia") or "inicial")],
        [chain_label, str(chain.get("cadeia") or "-")],
        ["Suporte / P(A-B-C)", f"{int(chain.get('suporte') or 0)} / {_pct(_finite(chain.get('probabilidade_conjunta')))}"],
        ["Função informada predominante", str(chain.get("funcao_predominante") or "não informada")],
        ["Cadeias temporais aceitas / estáveis", f"{accepted} / {stable}"],
    ]
    if chain_scope == "all":
        rows.insert(7, ["Cadeias incluídas", str(len(analysis.get("cadeias_completas") or []))])
    return rows


def _top_chain_rows(payload: dict[str, Any]) -> list[list[str]]:
    chains = [
        item
        for item in sorted((payload.get("analysis") or {}).get("cadeias_completas") or [], key=lambda item: int(item.get("suporte") or 0), reverse=True)
        if int(item.get("suporte") or 0) >= 8
    ]
    if str(payload.get("chain_scope") or "selected") == "selected":
        selected = payload.get("selected_chain") or {}
        chains = [selected] if selected and int(selected.get("suporte") or 0) >= 8 else []
    rows = []
    for item in chains:
        rows.append(
            [
                str(item.get("cadeia") or "-"),
                str(int(item.get("suporte") or 0)),
                _pct(_finite(item.get("probabilidade_conjunta"))),
                str(int(item.get("c1_leve") or 0)),
                str(int(item.get("c2_intenso") or 0)),
                _number_or_na(item.get("peso_medio_gravidade", item.get("indice_perigo"))),
            ]
        )
    return rows or [["Nenhuma cadeia atingiu suporte mínimo 8", "0", "0.0%", "0", "0", "-"]]


def _indicator_rows(payload: dict[str, Any]) -> list[list[str]]:
    return [
        ["Registros no recorte", "Total de episódios ABC que entraram nos cálculos após os filtros de ambiente e período."],
        ["C1 / C2", "C1 indica episódio leve, sem dano, lesão, sangramento ou direção a ponto vital. C2 indica ao menos um desses critérios de intensidade."],
        ["Frequência observada", "Proporção k/n do comportamento entre episódios ABC registrados no recorte; não é risco absoluto do próximo registro."],
        ["IC de Wilson", "Intervalo de confiança da proporção binomial bruta k/n, com numerador e denominador explícitos."],
        ["Qualidade da evidência", "Objeto auditável que considera amostra, sessões, dias, ausências, classes, validação temporal, calibração e validade externa."],
        ["Suporte (n)", "Número de registros completos em que a cadeia A-B-C foi observada."],
        ["P(A-B-C)", "Proporção dos registros do recorte que apresentaram a cadeia completa antecedente-comportamento-consequência."],
        ["P(B|A)", "Chance observada de o comportamento B ocorrer quando o antecedente A esteve presente."],
        ["P(C|A,B)", "Chance observada de a consequência C ocorrer depois do par antecedente A e comportamento B."],
        ["Lift", "Compara a frequência da cadeia com o esperado se seus elementos fossem independentes. Acima de 1 sugere associação maior que essa referência."],
        ["Peso médio de gravidade", "Média dos pesos internos C1=0,20 e C2=1,00 entre classificados. Não é probabilidade de dano nem escala validada."],
        ["Risco exploratório", "Combina frequência observada e peso médio configurado para priorização interna. Não representa diagnóstico nem causalidade."],
        ["Hipótese funcional registrada", "Hipótese mais frequentemente registrada para a cadeia. Não confirma função e não substitui avaliação funcional."],
        ["Temporais aceitas / estáveis", "Transições temporais aceitas pela revisão e quantas já possuem repetição, sessões e períodos suficientes para estabilidade."],
    ]


def _scope_label(payload: dict[str, Any]) -> str:
    if str(payload.get("chain_scope") or "selected") == "all":
        return "Todas as cadeias comportamentais do recorte"
    return "Uma cadeia comportamental específica"


def _chain_section_heading(payload: dict[str, Any]) -> str:
    if str(payload.get("chain_scope") or "selected") == "all":
        return "Cadeias A-B-C com suporte mínimo no corpo principal"
    return "Cadeia A-B-C selecionada"


def _temporal_rows(payload: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in (payload.get("temporal_data") or {}).get("candidates") or []:
        if item.get("rejection_reason") == "superseded_detection":
            continue
        status = str(item.get("validation_status") or "candidate")
        if status not in {"accepted", "candidate"}:
            continue
        chain = (
            f"{item.get('origin_behavior_code') or '?'} -> {item.get('from_consequence_code') or '?'} => "
            f"{item.get('to_antecedent_code') or '?'} -> {item.get('next_behavior_code') or '?'}"
        )
        rows.append([chain, str(item.get("environment") or "-"), str(item.get("delta_seconds") or "-"), "Aceita" if status == "accepted" else "Pendente"])
        if len(rows) >= 6:
            break
    return rows or [["Sem cadeia temporal disponível", "-", "-", "-"]]


def _methodology_lines(payload: dict[str, Any]) -> list[str]:
    prediction = payload.get("prediction") or {}
    mode = str(prediction.get("analysis_mode") or "descriptive")
    return [
        "Unidade descritiva: episódio ABC registrado. Ausência de registro não é convertida em ausência de comportamento.",
        "P(A|E) = n(A,E) / n(E); P(B|A,E) = n(A,B,E) / n(A,E); P(C|A,B,E) = n(A,B,C,E) / n(A,B,E).",
        "P(A,B,C|E) = n(A,B,C,E) / n(E). A identidade fatorada só é usada com a mesma unidade, ambiente, período e denominadores compatíveis.",
        "Opção frequencista: estimativa k/n e IC de Wilson para proporção binomial bruta. Opção bayesiana: prior Beta configurável e intervalo de credibilidade Beta.",
        "Como episódios da mesma sessão podem ser correlacionados, o relatório apresenta IC por bootstrap agrupado por sessão quando há pelo menos duas sessões identificáveis.",
        "Taxas por hora são calculadas somente com intervalos explicitamente observados e duração positiva; sem exposição, permanecem apenas contagens.",
        "Pesos configurados: C1=0,20 e C2=1,00. Gravidade esperada = soma dos pesos / classificados; não é probabilidade clínica.",
        "Risco exploratório = frequência observada × peso médio configurado de gravidade; índice interno não validado.",
        (
            "O módulo preditivo foi exibido porque apresentou alvo, horizonte, separação cronológica, baselines múltiplos, backtesting, métricas operacionais, calibração, drift e explicação individual."
            if mode == "predictive"
            else "O módulo preditivo foi bloqueado: o recorte não possui oportunidades negativas observáveis e validação temporal fora da amostra."
        ),
        "Transições temporais usam ID determinístico, gap não negativo e máximo configurável, deduplicação, agrupamento e critérios versionados de ocorrências, sessões e dias.",
        "Dataset, pipeline, taxonomia disponível e parâmetros da execução recebem versão ou hash reproduzível; validade externa continua explicitamente desconhecida.",
        f"Versão da metodologia: {(payload.get('analysis') or {}).get('methodology_version') or METHODOLOGY_VERSION}.",
    ]


def _period_label(payload: dict[str, Any]) -> str:
    series = pd.DataFrame((payload.get("analysis") or {}).get("serie_temporal") or [])
    if series.empty or "data" not in series:
        return "período não disponível"
    dates = pd.to_datetime(series["data"], errors="coerce").dropna()
    if dates.empty:
        return "período não disponível"
    return f"{dates.min().strftime('%d/%m/%Y')} a {dates.max().strftime('%d/%m/%Y')}"


def _configure_docx_header_footer(section) -> None:
    for header in (section.header, section.even_page_header, section.first_page_header):
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = "SELLAS PROJECT  |  ANÁLISE ABC V3"
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_paragraph.paragraph_format.space_after = Pt(0)
        for run in header_paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)

    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        prefix = footer_paragraph.add_run("Página ")
        prefix.font.name = "Arial"
        prefix.font.size = Pt(8)
        prefix.font.color.rgb = RGBColor(100, 100, 100)
        field_run = footer_paragraph.add_run()
        field_run.font.name = "Arial"
        field_run.font.size = Pt(8)
        field_run.font.color.rgb = RGBColor(100, 100, 100)
        field_begin = OxmlElement("w:fldChar")
        field_begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        field_separate = OxmlElement("w:fldChar")
        field_separate.set(qn("w:fldCharType"), "separate")
        field_result = OxmlElement("w:t")
        field_result.text = "1"
        field_end = OxmlElement("w:fldChar")
        field_end.set(qn("w:fldCharType"), "end")
        for element in (field_begin, instruction, field_separate, field_result, field_end):
            field_run._r.append(element)


def _configure_docx_styles(doc: Document) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.10
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    doc.styles["Heading 1"].paragraph_format.keep_with_next = True
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 2"].paragraph_format.keep_with_next = True
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(4)


def _add_docx_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    prevent_row_split: bool = True,
) -> None:
    if len(headers) != len(widths) or abs(sum(widths) - 6.5) > 0.01:
        raise ValueError("A geometria da tabela DOCX deve somar exatamente 6,5 polegadas.")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), "9360")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    table_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    for index, (cell, header) in enumerate(zip(table.rows[0].cells, headers, strict=False)):
        cell.text = header
        cell.width = Inches(widths[index])
        _set_cell_geometry(cell, round(widths[index] * 1440))
        _shade_cell(cell, "F2F4F7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row in rows:
        cells = table.add_row().cells
        if prevent_row_split:
            row_properties = table.rows[-1]._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            row_properties.append(cant_split)
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].width = Inches(widths[index])
            _set_cell_geometry(cells[index], round(widths[index] * 1440))
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)


def _set_cell_geometry(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _add_docx_ai_text(doc: Document, ai_result: Any) -> None:
    text = _clean_ai_text(ai_result)
    if not text:
        return
    doc.add_heading("Leitura funcional assistida já gerada", level=1)
    doc.add_paragraph(text)


def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SellasTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor(TEXT), alignment=TA_CENTER, spaceAfter=9))
    styles.add(ParagraphStyle(name="SellasH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=colors.HexColor("#405f3d"), spaceBefore=7, spaceAfter=5))
    styles.add(ParagraphStyle(name="SellasH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor("#405f3d"), spaceBefore=5, spaceAfter=3))
    styles.add(ParagraphStyle(name="SellasBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.8, leading=11.5, textColor=colors.HexColor(TEXT), spaceAfter=4))
    styles.add(ParagraphStyle(name="SellasWarning", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=8.8, leading=11.5, textColor=colors.HexColor("#803e37"), backColor=colors.HexColor("#f8e8e3"), borderPadding=6, spaceBefore=6))
    styles.add(ParagraphStyle(name="SellasFlow", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=11.5, leading=15, textColor=colors.HexColor("#405f3d"), backColor=colors.HexColor("#f3eadb"), borderPadding=8, alignment=TA_CENTER, spaceBefore=6, spaceAfter=6))
    styles.add(ParagraphStyle(name="SellasTable", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.7, leading=9.2, textColor=colors.HexColor(TEXT), spaceAfter=0))
    styles.add(ParagraphStyle(name="SellasTableHead", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=7.4, leading=8.7, textColor=colors.HexColor(TEXT), spaceAfter=0))
    return styles


def _draw_pdf_footer(canvas, document) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#d7cab5"))
    canvas.setLineWidth(0.4)
    canvas.line(1.35 * cm, 0.82 * cm, A4[0] - 1.35 * cm, 0.82 * cm)
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#6f675d"))
    canvas.drawString(1.35 * cm, 0.48 * cm, "Sellas Project - resumo ABC descritivo")
    canvas.drawRightString(A4[0] - 1.35 * cm, 0.48 * cm, f"Página {document.page}")
    canvas.restoreState()


def _pdf_table(headers: list[str], rows: list[list[str]], widths: list[float], styles) -> Table:
    wrapped = [
        [Paragraph(_escape(value), styles["SellasTableHead"]) for value in headers],
        *[[Paragraph(_escape(value), styles["SellasTable"]) for value in row] for row in rows],
    ]
    table = Table(wrapped, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#ede1ce")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9b897")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _figure_bytes(fig) -> io.BytesIO:
    output = io.BytesIO()
    fig.savefig(output, format="png", dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    output.seek(0)
    return output


def _clean_ai_text(ai_result: Any) -> str:
    if not isinstance(ai_result, dict):
        return ""
    text = str(ai_result.get("resposta") or "").strip()
    text = re.sub(r"[`#*_]+", "", text)
    text = re.sub(r"\s+", " ", text)
    return text[:3500]


def _finite(value: Any, default: float = 0.0) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return default
    return number if math.isfinite(number) else default


def _decimal_pt(value: Any, decimals: int = 1) -> str:
    if value is None or value == "":
        return "Não disponível"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "Não disponível"
    if not math.isfinite(number):
        return "Não disponível"
    return f"{number:.{decimals}f}".replace(".", ",")


def _pct(value: float) -> str:
    return f"{_finite(value) * 100:.1f}%".replace(".", ",")


def _pct_or_na(value: Any) -> str:
    if value is None or value == "":
        return "Não calculável"
    return f"{_finite(value) * 100:.1f}%".replace(".", ",")


def _number_or_na(value: Any) -> str:
    if value is None or value == "":
        return "Não calculável"
    return f"{_finite(value):.3f}".replace(".", ",")


def _short(value: Any, limit: int) -> str:
    text = str(value or "-")
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def _wrap_chart_label(value: Any, width: int) -> str:
    return "\n".join(textwrap.wrap(str(value or "-"), width=max(8, width), break_long_words=False, break_on_hyphens=False))


def _escape(value: Any) -> str:
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
