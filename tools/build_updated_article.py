from __future__ import annotations

import math
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "artigo_previsao_comportamental_bHave_integrado_mvp.docx"
OUTPUT = ROOT / "docs" / "artigo_previsao_comportamental_bHave_integrado_mvp_logica.docx"
FIG_DIR = ROOT / "docs" / "rendered_artigo_bhave_word" / "logic_figures"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "DDEFE5"
GOLD = "FFF2CC"
RED = "FCE4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D9E2F3") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_table(table, header_fill: str = LIGHT_BLUE) -> None:
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.style = "CaptionCustom" if "CaptionCustom" in [s.name for s in doc.styles] else "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_equation_box(doc: Document, label: str, equations: list[str], explanation: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    set_cell_shading(table.cell(0, 0), LIGHT_BLUE)
    set_cell_shading(table.cell(1, 0), "FFFFFF")
    for cell in table._cells:
        set_cell_border(cell)

    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)

    body = table.cell(1, 0)
    for idx, equation in enumerate(equations):
        paragraph = body.paragraphs[0] if idx == 0 else body.add_paragraph()
        run = paragraph.add_run(equation)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9.5)
    paragraph = body.add_paragraph()
    paragraph.add_run(explanation)
    paragraph.paragraph_format.space_before = Pt(4)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "B7C9DF")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    paragraph.add_run(" " + body)
    doc.add_paragraph()


def draw_flow_diagram(path: Path) -> None:
    boxes = [
        ("bHave API", "dados estruturados\ne anotacoes"),
        ("Contrato canonico", "normalizacao\nschema unico"),
        ("Feature store", "historico, janelas,\ntexto e contexto"),
        ("Validacao temporal", "paciente + periodo\nsem vazamento"),
        ("Modelo calibrado", "probabilidade\n0 a 1"),
        ("Painel clinico", "risco, fatores,\nalerta de limite"),
    ]
    fig, ax = plt.subplots(figsize=(11, 3.2), dpi=220)
    ax.axis("off")
    for i, (title, subtitle) in enumerate(boxes):
        x = i * 1.75
        rect = plt.Rectangle((x, 0.55), 1.35, 1.1, facecolor="#EEF6F2" if i in (2, 4) else "#E8EEF5", edgecolor="#2E74B5", linewidth=1.8)
        ax.add_patch(rect)
        ax.text(x + 0.675, 1.28, title, ha="center", va="center", fontsize=9.5, fontweight="bold", color="#0B2545")
        ax.text(x + 0.675, 0.91, subtitle, ha="center", va="center", fontsize=8.2, color="#3B4A5A")
        if i < len(boxes) - 1:
            ax.annotate("", xy=(x + 1.62, 1.1), xytext=(x + 1.36, 1.1), arrowprops=dict(arrowstyle="->", color="#1F4D78", lw=1.6))
    ax.set_xlim(-0.15, 10.1)
    ax.set_ylim(0.25, 1.95)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def draw_skill_forecast(path: Path) -> None:
    months = [1, 2, 3, 4, 5, 6]
    observed = [70, 20, 100]
    weights = [0.2, 0.3, 0.5]
    xbar = sum(w * m for w, m in zip(weights, months[:3])) / sum(weights)
    ybar = sum(w * y for w, y in zip(weights, observed)) / sum(weights)
    numerator = sum(w * (m - xbar) * (y - ybar) for w, m, y in zip(weights, months[:3], observed))
    denominator = sum(w * (m - xbar) ** 2 for w, m in zip(weights, months[:3]))
    slope = numerator / denominator
    projection = [max(0, min(100, observed[-1] + slope * h)) for h in range(1, 4)]

    fig, ax = plt.subplots(figsize=(7.3, 4.2), dpi=220)
    ax.plot(months[:3], observed, marker="o", linewidth=2.4, color="#2E74B5", label="observado")
    ax.plot(months[3:], projection, marker="o", linewidth=2.4, color="#7A5A00", linestyle="--", label="projetado")
    ax.fill_between(months[3:], [max(0, y - 12) for y in projection], [min(100, y + 12) for y in projection], color="#FFF2CC", alpha=0.8, label="faixa esperada")
    ax.axhline(80, color="#15803D", linewidth=1.5, linestyle=":", label="criterio clinico 80%")
    ax.set_title("Previsao de independencia por mes", fontsize=13, fontweight="bold", color="#0B2545")
    ax.set_xlabel("Mes")
    ax.set_ylabel("Independencia media (%)")
    ax.set_ylim(0, 110)
    ax.grid(True, alpha=0.25)
    ax.legend(loc="lower right", fontsize=8)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def draw_behavior_forecast(path: Path) -> None:
    sessions = list(range(1, 13))
    counts = [0, 1, 0, 3, 2, 2, 1, 4, 3, 2, 1, 0]
    alpha = 0.3
    risk = []
    current = 0.0
    for y in [1 if c > 0 else 0 for c in counts]:
        current = alpha * y + (1 - alpha) * current
        risk.append(current)
    probability = [0.18 + 0.62 * r + 0.03 * math.log1p(c) for r, c in zip(risk, counts)]

    fig, ax1 = plt.subplots(figsize=(7.3, 4.2), dpi=220)
    ax1.bar(sessions, counts, color="#D9EAF7", edgecolor="#2E74B5", label="contagem")
    ax1.set_xlabel("Sessao")
    ax1.set_ylabel("Contagem", color="#1F4D78")
    ax1.tick_params(axis="y", labelcolor="#1F4D78")
    ax2 = ax1.twinx()
    ax2.plot(sessions, probability, color="#9B1C1C", marker="o", linewidth=2.2, label="p calibrada")
    ax2.axhspan(0.7, 1.0, color="#FCE4D6", alpha=0.5)
    ax2.axhspan(0.3, 0.7, color="#FFF2CC", alpha=0.35)
    ax2.set_ylabel("Probabilidade prevista", color="#9B1C1C")
    ax2.set_ylim(0, 1)
    ax2.tick_params(axis="y", labelcolor="#9B1C1C")
    ax1.set_title("Previsao de comportamento-problema", fontsize=13, fontweight="bold", color="#0B2545")
    ax1.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def draw_notes_logic(path: Path) -> None:
    boxes = [
        ("Texto bruto", "\"chorou 3 vezes\""),
        ("Normalizacao", "lowercase\nsem acentos"),
        ("Regras NLP", "comportamento\nnegacao\ncontexto"),
        ("Revisao humana", "corrige ambiguos"),
        ("Features", "ocorrencia\nfreq.\nintensidade"),
    ]
    fig, ax = plt.subplots(figsize=(10, 3.5), dpi=220)
    ax.axis("off")
    for i, (title, text) in enumerate(boxes):
        x = i * 2
        fill = ["#F2F4F7", "#E8EEF5", "#FFF2CC", "#FCE4D6", "#DDEFE5"][i]
        rect = plt.Rectangle((x, 0.65), 1.45, 1.25, facecolor=fill, edgecolor="#6B7280", linewidth=1.4)
        ax.add_patch(rect)
        ax.text(x + 0.725, 1.48, title, ha="center", va="center", fontsize=9.5, fontweight="bold", color="#0B2545")
        ax.text(x + 0.725, 1.05, text, ha="center", va="center", fontsize=8.3, color="#374151")
        if i < len(boxes) - 1:
            ax.annotate("", xy=(x + 1.85, 1.25), xytext=(x + 1.46, 1.25), arrowprops=dict(arrowstyle="->", color="#1F4D78", lw=1.5))
    ax.set_xlim(-0.1, 9.5)
    ax.set_ylim(0.35, 2.2)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)


def add_logic_section(doc: Document) -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    flow = FIG_DIR / "figura_fluxo_logico.png"
    skill = FIG_DIR / "figura_previsao_habilidades.png"
    behavior = FIG_DIR / "figura_previsao_comportamento.png"
    notes = FIG_DIR / "figura_extracao_anotacoes.png"
    draw_flow_diagram(flow)
    draw_skill_forecast(skill)
    draw_behavior_forecast(behavior)
    draw_notes_logic(notes)

    doc.add_page_break()
    doc.add_heading("PARTE III - Logica matematica, evolucao clinica e explicabilidade operacional", level=1)
    add_callout(
        doc,
        "Resumo da atualizacao.",
        "Esta parte integra a revisao academica ao MVP implementado: mostra como os dados da bHave entram no contrato canonico, como as evolucoes/anotacoes sao transformadas em variaveis, quais equacoes sustentam as previsoes e quais salvaguardas impedem uso clinico indevido.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("1. Visao geral da logica do sistema", level=2)
    p = doc.add_paragraph()
    p.add_run("A logica do aplicativo foi desenhada como uma cadeia temporal. ").bold = True
    p.add_run(
        "Cada sessao do paciente gera registros estruturados e, quando disponivel, texto clinico. "
        "Esses dados sao normalizados, convertidos em features e usados para estimar uma probabilidade futura. "
        "A probabilidade nao e uma decisao automatica; ela e um sinal analitico para apoiar a leitura do analista do comportamento."
    )
    doc.add_picture(str(flow), width=Inches(6.35))
    add_caption(doc, "Figura 1. Fluxo logico do dado ate a previsao calibrada.")

    table = doc.add_table(rows=1, cols=4)
    headers = ["Parte", "Pergunta que responde", "Entrada", "Saida"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    rows = [
        ("Adapter bHave", "O dado veio no contrato esperado?", "API real, mock ou arquivo normalizado", "Schema canonico auditavel"),
        ("Feature engineering", "O que ja era conhecido antes da previsao?", "Historico ate t", "Vetor X_t sem vazamento"),
        ("Extracao de anotacoes", "O texto clinico contem sinal comportamental?", "Evolucoes e notas", "Features textuais revisaveis"),
        ("Modelo preditivo", "Qual a chance de ocorrer no horizonte H?", "X_t", "p(Y_t,h = 1 | X_t)"),
        ("Calibracao", "70% significa aproximadamente 70%?", "Probabilidade bruta + validacao", "Probabilidade calibrada"),
        ("Governanca", "Pode ser usado com seguranca?", "Metricas, drift, LGPD, auditoria", "Status de uso permitido ou bloqueado"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table)

    doc.add_heading("2. Variavel-alvo e horizonte de previsao", level=2)
    doc.add_paragraph(
        "Para comportamento-problema, a variavel-alvo e binaria. Em cada sessao t, o sistema pergunta se o comportamento X ocorrera em uma janela futura h. "
        "No MVP, h costuma representar a proxima sessao, mas o contrato aceita horizontes configuraveis."
    )
    add_equation_box(
        doc,
        "Equacao 1 - alvo temporal",
        [
            "Y_{t,h} = 1 se o comportamento X ocorrer na janela futura h",
            "Y_{t,h} = 0 se o comportamento X nao ocorrer na janela futura h",
            "p_{t,h} = P(Y_{t,h}=1 | X_t)",
        ],
        "O vetor X_t contem somente informacoes conhecidas ate a sessao t. Essa regra evita vazamento de informacao futura.",
    )

    doc.add_heading("3. Previsao de habilidades e objetivos", level=2)
    doc.add_paragraph(
        "Para habilidades e objetivos, o foco nao e apenas classificar risco. O sistema estima a trajetoria de independencia ao longo dos meses. "
        "Se um paciente apresenta 70% de independencia no mes 1, 20% no mes 2 e 100% no mes 3, a aplicacao calcula uma tendencia suavizada, limita a previsao entre 0% e 100% e informa a probabilidade de atingir o criterio clinico."
    )
    doc.add_picture(str(skill), width=Inches(6.1))
    add_caption(doc, "Figura 2. Exemplo de previsao mensal de independencia a partir de tres meses observados.")
    add_equation_box(
        doc,
        "Equacao 2 - tendencia ponderada de independencia",
        [
            "I_m = independencia media observada no mes m",
            "b = soma[w_m (m - m_barra)(I_m - I_barra)] / soma[w_m (m - m_barra)^2]",
            "I_hat_{m+h} = clip(I_m + h*b, 0, 100)",
            "P(meta) = sigmoid((I_hat_{m+h} - theta) / sigma)",
        ],
        "Pesos maiores podem ser dados aos meses recentes. theta representa o criterio clinico, por exemplo 80% de independencia; sigma controla quao abrupta e a transicao de probabilidade.",
    )
    add_callout(
        doc,
        "Leitura clinica.",
        "Uma queda de 70% para 20% e depois subida para 100% nao deve ser interpretada como linha reta simples. O sistema exibe tendencia, variabilidade e criterio, permitindo que o profissional avalie se a oscilacao reflete aprendizagem, mudanca de demanda, prompt, ambiente ou registro inconsistente.",
        fill=GREEN,
    )

    doc.add_heading("4. Previsao de comportamento-problema", level=2)
    doc.add_paragraph(
        "Para comportamento-problema, o MVP combina baseline historico, risco recente, janelas moveis, contexto, antecedentes, consequencias e features textuais. "
        "A previsao final e uma probabilidade calibrada, acompanhada de fatores explicativos e classificacao de risco."
    )
    doc.add_picture(str(behavior), width=Inches(6.1))
    add_caption(doc, "Figura 3. Exemplo de contagem observada e probabilidade calibrada por sessao.")
    add_equation_box(
        doc,
        "Equacao 3 - risco recente exponencial",
        [
            "R_t = alpha*y_t + (1-alpha)*R_{t-1}",
            "F_recente(t,N) = soma dos eventos nas ultimas N sessoes",
            "p_baseline = sessoes com ocorrencia / sessoes observadas",
        ],
        "O parametro alpha controla quanto o evento mais recente pesa na avaliacao. Valores maiores tornam o risco mais sensivel ao ultimo registro.",
    )
    add_equation_box(
        doc,
        "Equacao 4 - modelo probabilistico",
        [
            "z = beta_0 + beta_1*x_1 + ... + beta_n*x_n",
            "p_bruto = 1 / (1 + exp(-z))",
            "p_calibrado = sigmoid(a*p_bruto + b)",
        ],
        "A regressao logistica fornece uma probabilidade bruta. A calibracao ajusta essa saida para que a interpretacao percentual seja mais confiavel.",
    )

    doc.add_heading("5. Como as evolucoes e anotacoes entram no modelo", level=2)
    doc.add_paragraph(
        "As evolucoes clinicas e anotacoes de atendimento foram adicionadas como fonte complementar. "
        "O sistema nao assume que todo texto e verdade operacional final: ele extrai sinais, marca ambiguidade e permite correcao humana."
    )
    doc.add_picture(str(notes), width=Inches(6.2))
    add_caption(doc, "Figura 4. Pipeline de extracao de anotacoes clinicas para features preditivas.")
    add_equation_box(
        doc,
        "Equacao 5 - transformacao textual em features",
        [
            "s_limpo = normalizar(texto_bruto)",
            "o_b in {0, 1, nulo} para cada comportamento b",
            "X_t = [features_estruturadas_t, features_textuais_t]",
            "prioridade = correcao humana > dado estruturado > extracao automatica > desconhecido",
        ],
        "Negacoes reduzem ocorrencia; termos ambiguos reduzem confianca e acionam revisao humana. Correcoes humanas confirmadas passam a ter prioridade.",
    )
    add_bullets(
        doc,
        [
            "Termos como 'nao houve', 'sem' e 'ausencia de' sao tratados como negacao.",
            "Termos como 'pareceu', 'possivel', 'quase' e 'tentou' indicam incerteza.",
            "Frequencia, intensidade, duracao, antecedente, consequencia e contexto viram colunas numericas ou binarias.",
            "A revisao humana corrige falsos positivos, falsos negativos e ambiguidades clinicas.",
        ],
    )

    doc.add_heading("6. Validacao temporal, calibracao e monitoramento", level=2)
    doc.add_paragraph(
        "A validacao precisa respeitar o tempo e a identidade do paciente. Misturar sessoes futuras no treino ou testar no mesmo periodo usado para escolher limiar gera desempenho artificialmente alto. "
        "Por isso, a avaliacao e feita com corte temporal e monitoramento continuo."
    )
    add_equation_box(
        doc,
        "Equacao 6 - Brier Score e drift",
        [
            "BS = media((p_i - y_i)^2)",
            "PSI = soma_i (p_i - q_i) * ln(p_i / q_i)",
        ],
        "O Brier Score mede qualidade probabilistica. O PSI compara distribuicoes entre referencia e producao para sinalizar drift de dados.",
    )
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Controle", "Logica", "Motivo clinico"]):
        table.cell(0, idx).text = header
    rows = [
        ("Split temporal", "Treino em sessoes antigas; teste em sessoes recentes", "Simula previsao real"),
        ("Agrupamento por paciente", "Metricas por paciente e periodo", "Evita esconder heterogeneidade"),
        ("Calibracao", "Brier, curva de calibracao e Platt/isotonic quando aplicavel", "Torna o percentual interpretavel"),
        ("Drift", "Comparacao de distribuicao e desempenho", "Detecta mudanca de ambiente, registro ou perfil"),
        ("Abstencao", "Bloqueio quando dado esta fora do dominio", "Evita alerta com baixa confianca"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table, header_fill=GOLD)

    doc.add_heading("7. Interpretacao dos resultados no painel", level=2)
    doc.add_paragraph(
        "A leitura recomendada segue quatro perguntas: a probabilidade esta acima do baseline? O risco recente explica a mudanca? Existem anotacoes textuais confirmando o padrao? As metricas e o drift permitem confiar no modelo neste periodo?"
    )
    add_bullets(
        doc,
        [
            "Risco baixo: p < 0,30; normalmente monitoramento e revisao de rotina.",
            "Risco moderado: 0,30 <= p < 0,70; revisar contexto, antecedentes, consequencias e tendencia.",
            "Risco alto: p >= 0,70; exige leitura clinica cuidadosa, mas nao autoriza intervencao automatica.",
            "Probabilidade alta com baixa confianca textual ou drift ativo deve ser tratada como alerta tecnico, nao como conclusao clinica.",
        ],
    )

    doc.add_heading("8. Limites, LGPD e uso real", level=2)
    add_callout(
        doc,
        "Limite regulatorio.",
        "O artigo e o MVP descrevem apoio a decisao e pesquisa aplicada. Antes de uso real, e necessario revisar LGPD, governanca clinica, auditoria, consentimento/base legal, seguranca, vieses e possivel enquadramento regulatorio como software em saude.",
        fill=RED,
    )
    doc.add_paragraph(
        "A previsao nao substitui avaliacao funcional, julgamento profissional, analise de contingencias nem decisao do analista do comportamento. "
        "O modelo descreve padroes observados nos dados disponiveis; ele nao prova causa, funcao comportamental ou eficacia de intervencao."
    )

    doc.add_heading("9. Atualizacao do estado de verificacao", level=2)
    add_bullets(
        doc,
        [
            "Testes automatizados do modulo bHave executados com sucesso: 23 passed.",
            "Foram adicionadas features textuais derivadas de evolucoes/anotacoes de atendimento.",
            "Foram adicionados endpoints de atendimentos, extracao, revisao humana e features from-notes.",
            "Foi registrado documento de fix em docs/fixes/13_07_26_fix16.md.",
        ],
    )


def update_existing_verification_text(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if "Testes automatizados do modulo bHave executados com sucesso: 13 passed" in paragraph.text:
            paragraph.text = "Testes automatizados do modulo bHave executados com sucesso: 23 passed."


def set_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()
        if not paragraph.text.strip():
            paragraph.text = "bHave | artigo tecnico-academico atualizado"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(89, 89, 89)


def main() -> None:
    doc = Document(SOURCE)
    update_existing_verification_text(doc)
    add_logic_section(doc)
    set_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
