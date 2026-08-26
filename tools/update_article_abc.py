from __future__ import annotations

from datetime import datetime
from pathlib import Path

import matplotlib
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


matplotlib.use("Agg")
import matplotlib.pyplot as plt


PROJECT_DIR = Path(__file__).resolve().parents[1]
ARTICLE_PATH = PROJECT_DIR / "docs" / "artigo_previsao_comportamental_bHave_integrado_mvp_logica.docx"
FIGURE_DIR = PROJECT_DIR / "docs" / "figures"
CHAIN_FIGURE = FIGURE_DIR / "figura_5_cadeia_abc_probabilidade.png"
HEATMAP_FIGURE = FIGURE_DIR / "figura_6_heatmap_previsao_abc.png"

GREEN = "405F3D"
LIGHT_GREEN = "DCE8DF"
LIGHT_GOLD = "F2E7D5"
LIGHT_BLUE = "DDE6EF"
ROSE = "EEDBD7"
TEXT = "241A12"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_equation_box(doc: Document, title: str, equations: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GOLD)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(TEXT)
    for equation in equations:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(equation)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caption(doc: Document, text: str) -> None:
    style = "CaptionCustom" if "CaptionCustom" in [style.name for style in doc.styles] else None
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def draw_chain_figure() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(11.5, 4.6), dpi=180)
    fig.patch.set_facecolor("#fffaf2")
    ax.set_facecolor("#fffaf2")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis("off")

    boxes = [
        (0.35, 2.05, 2.25, 1.15, "Contexto E", "ambiente + horário\ncom segundos", "#dde6ef"),
        (3.15, 2.05, 2.15, 1.15, "A", "antecedente\nobservável", "#f2e7d5"),
        (5.85, 2.05, 2.15, 1.15, "B", "comportamento\noperacionalizado", "#dce8df"),
        (8.55, 2.05, 2.15, 1.15, "C", "consequência\nobservável", "#eedbd7"),
    ]
    for x, y, width, height, title, subtitle, color in boxes:
        patch = plt.matplotlib.patches.FancyBboxPatch(
            (x, y),
            width,
            height,
            boxstyle="round,pad=0.03,rounding_size=0.08",
            facecolor=color,
            edgecolor="#6b6259",
            linewidth=1.2,
        )
        ax.add_patch(patch)
        ax.text(x + width / 2, y + 0.73, title, ha="center", va="center", fontsize=13, fontweight="bold", color="#241a12")
        ax.text(x + width / 2, y + 0.34, subtitle, ha="center", va="center", fontsize=9.2, color="#4f463d")

    for start, end, label in ((2.60, 3.15, "P(A|E)"), (5.30, 5.85, "P(B|A,E)"), (8.00, 8.55, "P(C|A,B,E)")):
        ax.annotate("", xy=(end, 2.63), xytext=(start, 2.63), arrowprops=dict(arrowstyle="->", color="#405f3d", lw=2.0))
        ax.text((start + end) / 2, 2.93, label, ha="center", va="bottom", fontsize=8.5, color="#405f3d", fontweight="bold")

    ax.text(
        6,
        1.18,
        "P(A,B,C | E) = P(A | E) × P(B | A,E) × P(C | A,B,E)",
        ha="center",
        va="center",
        fontsize=12,
        fontweight="bold",
        color="#405f3d",
    )
    ax.text(
        6,
        0.55,
        "A cadeia é uma associação temporal observada; o cálculo não demonstra causalidade nem função comportamental.",
        ha="center",
        va="center",
        fontsize=9.5,
        color="#6b6259",
    )
    fig.tight_layout(pad=0.6)
    fig.savefig(CHAIN_FIGURE, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def draw_heatmap_figure() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    fig, axes = plt.subplots(1, 2, figsize=(11.5, 5.6), dpi=180, gridspec_kw={"width_ratios": [1.35, 1]})
    fig.patch.set_facecolor("#fffaf2")
    matrix = np.array([[0.72, 0.18, 0.08], [0.20, 0.64, 0.16], [0.08, 0.18, 0.76]])
    ax = axes[0]
    image = ax.imshow(matrix, cmap="RdYlGn_r", vmin=0, vmax=1, aspect="auto")
    ax.set_xticks(range(3), ["Sala de terapia", "Sala de aula", "Casa"], fontsize=8.5)
    ax.set_yticks(range(3), ["Demanda → Choro → Pausa", "Espera → Grito → Atenção", "Transição → Fuga → Pausa"], fontsize=8.3)
    ax.set_title("Heatmap H(c,e) = n(c,e) / n(e)", fontsize=11.5, fontweight="bold", color="#241a12", pad=12)
    for row in range(matrix.shape[0]):
        for column in range(matrix.shape[1]):
            ax.text(column, row, f"{matrix[row, column] * 100:.0f}%", ha="center", va="center", fontsize=9, fontweight="bold", color="#241a12")
    colorbar = fig.colorbar(image, ax=ax, fraction=0.045, pad=0.04)
    colorbar.set_label("P(cadeia | ambiente)", fontsize=8.5)

    ax = axes[1]
    components = ["Baseline", "Recência", "Antecedente", "Ambiente", "A + ambiente"]
    weights = [0.35, 0.15, 0.20, 0.15, 0.15]
    colors = ["#6f86a4", "#c17c74", "#b58a55", "#7d9b76", "#6f9b96"]
    bars = ax.barh(components[::-1], weights[::-1], color=colors[::-1], edgecolor="#6b6259", linewidth=0.7)
    ax.set_xlim(0, 0.40)
    ax.set_xlabel("Peso antes da normalização", fontsize=9, labelpad=8)
    ax.set_title("Composição da previsão", fontsize=11.5, fontweight="bold", color="#241a12", pad=12)
    ax.grid(axis="x", color="#ded0b8", linewidth=0.7, alpha=0.8)
    ax.set_axisbelow(True)
    for bar, value in zip(bars, weights[::-1]):
        ax.text(value + 0.008, bar.get_y() + bar.get_height() / 2, f"{value:.2f}", va="center", fontsize=9)
    for axis in axes:
        axis.set_facecolor("#fffaf2")
    fig.tight_layout(rect=[0, 0.10, 1, 1], pad=1.2)
    fig.text(
        0.77,
        0.035,
        "p_próximo = Σ(wⱼ · pⱼ) / Σwⱼ",
        ha="center",
        va="center",
        fontsize=10.5,
        fontweight="bold",
        color="#405f3d",
    )
    fig.savefig(HEATMAP_FIGURE, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def build_article_update() -> None:
    if not ARTICLE_PATH.exists():
        raise FileNotFoundError(ARTICLE_PATH)
    draw_chain_figure()
    draw_heatmap_figure()

    doc = Document(ARTICLE_PATH)
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("PARTE IV - Registro ABC fechado, previsão contextual e apoio à avaliação funcional", level=1)
    lead = doc.add_paragraph()
    lead_run = lead.add_run(
        "Esta parte documenta a implementação adicionada ao Sellas Project em 14 de julho de 2026. "
        "O módulo registra cadeias Antecedente-Comportamento-Consequência com ambiente e precisão de segundos, "
        "calcula associações condicionais, apresenta um heatmap contextual e estima a probabilidade exploratória do próximo comportamento."
    )
    lead_run.bold = True

    doc.add_heading("1. Escopo operacional e separação no painel", level=2)
    doc.add_paragraph(
        "O painel passou a ter três áreas mutuamente exclusivas: previsão de habilidades/objetivos, previsão comportamental agregada e previsão ABC fechada. "
        "Essa separação evita misturar séries mensais de desempenho com acontecimentos pontuais de contingência. A área ABC aceita categorias fechadas já existentes "
        "e permite digitar novos antecedentes, comportamentos, consequências e ambientes. A nova opção só é persistida quando o acontecimento é efetivamente adicionado."
    )
    add_bullet(doc, "Cada acontecimento contém paciente, data, hora, segundo, ambiente, antecedente, comportamento e consequência.")
    add_bullet(doc, "Adicionar ou remover recalcula a análise e regenera o Excel do paciente; a remoção permanece no log de auditoria.")
    add_bullet(doc, "Os registros continuam pseudonimizados pelo token do paciente e são vinculados às tabelas ABC versionadas.")

    doc.add_heading("2. Estrutura temporal e unidade de análise", level=2)
    doc.add_paragraph(
        "A unidade analítica é um acontecimento ABC registrado em um instante local t, no fuso America/Sao_Paulo. O segundo é armazenado no timestamp e exibido como HH:MM:SS. "
        "Isso permite distinguir duas cadeias que ocorreram no mesmo minuto, mas em segundos diferentes. O intervalo técnico continua delimitado para compatibilidade com o esquema de observação, "
        "enquanto o início do intervalo preserva o instante informado pelo usuário."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Campo", "Símbolo", "Tipo", "Uso analítico"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, GREEN)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    rows = [
        ("Ambiente", "E", "categórico", "filtro, estratificação e heatmap"),
        ("Antecedente", "A", "categórico", "exposição contextual observada"),
        ("Comportamento", "B", "categórico", "evento a descrever ou prever"),
        ("Consequência", "C", "categórico", "evento observado após A-B"),
        ("Data e hora", "t", "timestamp", "ordenação e linha do tempo por segundo"),
    ]
    fills = [LIGHT_GREEN, "FFFFFF"]
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cells[column_index].text = value
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[column_index])
            set_cell_shading(cells[column_index], fills[row_index % 2])

    doc.add_heading("3. Probabilidade da cadeia A-B-C", level=2)
    doc.add_paragraph(
        "Para um ambiente E, a frequência da cadeia completa é decomposta pela regra do produto. Essa decomposição permite localizar em qual transição a cadeia se concentra: "
        "na presença do antecedente, na ocorrência do comportamento após o antecedente ou na consequência após o par A-B."
    )
    add_equation_box(
        doc,
        "Equações da cadeia completa",
        [
            "P(B | A,E) = n(A,B,E) / n(A,E)",
            "P(C | A,B,E) = n(A,B,C,E) / n(A,B,E)",
            "P(A,B,C | E) = P(A | E) × P(B | A,E) × P(C | A,B,E)",
            "Lift conjunto = P(A,B,C) / [P(A) × P(B) × P(C)]",
        ],
    )
    doc.add_paragraph(
        "O suporte n(A,B,C,E) é sempre exibido junto da probabilidade. Uma probabilidade alta com suporte muito pequeno é instável e não deve receber o mesmo peso clínico de um padrão repetido. "
        "O lift compara a frequência conjunta observada à frequência esperada sob independência marginal; lift maior que 1 indica concentração estatística, não causalidade."
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(CHAIN_FIGURE), width=Inches(6.45))
    add_caption(doc, "Figura 5. Decomposição da probabilidade da cadeia ABC condicionada ao ambiente e ao instante observado.")

    doc.add_heading("4. Heatmap por ambiente e leitura temporal", level=2)
    doc.add_paragraph(
        "O heatmap compara cadeias completas entre ambientes. Para cada cadeia c e ambiente e, a célula apresenta a proporção de registros daquele ambiente que pertencem à cadeia. "
        "A linha do tempo usa o timestamp completo e mantém segundos no eixo e no hover, enquanto o gráfico diário usa datas categóricas para impedir a escala incorreta de microssegundos observada anteriormente."
    )
    add_equation_box(doc, "Valor de cada célula do heatmap", ["H(c,e) = n(c,e) / n(e)"])
    add_bullet(doc, "Comparar linhas identifica em quais ambientes uma mesma cadeia se concentra.")
    add_bullet(doc, "Comparar colunas mostra quais cadeias representam maior parcela dos registros de cada ambiente.")
    add_bullet(doc, "Ausência de célula não significa ausência real do comportamento; pode significar ausência de observação ou oportunidade.")

    doc.add_heading("5. Previsão exploratória do próximo comportamento", level=2)
    doc.add_paragraph(
        "A previsão ABC não usa uma classe causal. Ela estima a chance de o comportamento selecionado ser o próximo comportamento registrado, dadas as evidências disponíveis. "
        "Para reduzir extremos quando a amostra é pequena, cada proporção recebe suavização Beta(1,1), equivalente à correção de Laplace. A recência é representada por uma média móvel exponencial calculada sobre até 20 acontecimentos, com alfa igual a 0,30."
    )
    add_equation_box(
        doc,
        "Suavização, recência e composição",
        [
            "p_suavizada = (k + 1) / (n + 2)",
            "r_t = α·y_t + (1 - α)·r_(t-1), com α = 0,30 e y_t ∈ {0,1}",
            "p_próximo = Σ(w_j·p_j) / Σw_j",
            "w = {baseline: 0,35; recência: 0,15; antecedente: 0,20; ambiente: 0,15; A+ambiente: 0,15}",
        ],
    )
    doc.add_paragraph(
        "Quando antecedente ou ambiente não são selecionados, os componentes indisponíveis são retirados e os pesos restantes são normalizados para somar 1. "
        "O painel mostra baseline, probabilidade contextual, tamanho da amostra, qualidade descritiva da evidência e intervalo de Wilson de 95% para a proporção contextual."
    )
    add_equation_box(
        doc,
        "Intervalo de Wilson para k ocorrências em n observações",
        [
            "centro = [p̂ + z²/(2n)] / [1 + z²/n]",
            "margem = z·√[p̂(1-p̂)/n + z²/(4n²)] / [1 + z²/n]",
            "IC95% = centro ± margem, com z = 1,96",
        ],
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(HEATMAP_FIGURE), width=Inches(6.45))
    add_caption(doc, "Figura 6. Exemplo didático do heatmap por ambiente e dos componentes ponderados da previsão ABC.")

    doc.add_heading("6. Expansão para análise ou avaliação funcional", level=2)
    doc.add_paragraph(
        "A expansão funcional do painel é uma camada exploratória. Ela organiza a cadeia selecionada por ambiente, mostra repetição, suporte, consequência predominante e concentração temporal. "
        "Esses resultados podem apoiar a formulação de hipóteses e a decisão sobre quais situações observar com maior rigor, mas não constituem avaliação funcional concluída."
    )
    add_bullet(doc, "Definir o comportamento em termos observáveis, mensuráveis e replicáveis.")
    add_bullet(doc, "Registrar oportunidades, períodos não observados e exposição ao antecedente para evitar denominadores enganosos.")
    add_bullet(doc, "Medir frequência, duração, latência ou intensidade quando a pergunta clínica exigir.")
    add_bullet(doc, "Avaliar concordância entre observadores e revisar categorias personalizadas antes de agregá-las.")
    add_bullet(doc, "Combinar a análise descritiva com entrevista, observação direta e procedimentos funcionais apropriados conduzidos por profissional habilitado.")

    doc.add_heading("7. Persistência, auditoria e Excel por paciente", level=2)
    doc.add_paragraph(
        "Cada adição salva uma sessão, um intervalo e os eventos fechados no Supabase. Para as três categorias escolhidas, ocorreu=true; para as demais categorias ativas, ocorreu=false. "
        "O log de ações preserva adições, remoções e criação de categorias. O Excel contém resumo, acontecimentos ativos, log de ações e catálogo de categorias. Data e hora do acontecimento são separadas nos registros ativos, "
        "e o log mantém a hora da ação e a hora do acontecimento com segundos, normalizadas para America/Sao_Paulo."
    )

    doc.add_heading("8. Regras de interpretação e segurança clínica", level=2)
    add_bullet(doc, "Probabilidade é uma estimativa baseada no que foi registrado, não uma certeza sobre o que ocorrerá.")
    add_bullet(doc, "P(B|A,E) não equivale a dizer que A causou B.")
    add_bullet(doc, "P(C|A,B,E) não prova que C mantém B; descreve apenas a sequência observada.")
    add_bullet(doc, "Heatmap e lift dependem da qualidade dos denominadores, da consistência das categorias e da cobertura de observação.")
    add_bullet(doc, "Resultados com amostra inicial devem ser tratados como sinal para coleta adicional, não como base única de intervenção.")
    doc.add_paragraph(
        "Antes de uso real, permanecem necessários revisão LGPD, governança de acesso, rastreabilidade de versões, validação temporal por paciente e período, monitoramento de drift, governança clínica e avaliação de possível enquadramento regulatório."
    )

    doc.add_heading("9. Estado de verificação desta atualização", level=2)
    add_bullet(doc, "A terceira área de previsão foi validada no navegador com o paciente PACIENTE TESTE.")
    add_bullet(doc, "A digitação de nova opção foi validada sem persistir categoria temporária.")
    add_bullet(doc, "O campo de segundo, o filtro por ambiente, a linha do tempo, o heatmap, a cadeia completa e a previsão foram renderizados sem traceback.")
    add_bullet(doc, "O Excel foi regenerado e revisado nas quatro planilhas, com horário local e segundos.")
    add_bullet(doc, "Os testes focados do serviço ABC foram executados com sucesso: 4 passed.")
    add_bullet(doc, "As alterações estão documentadas em docs/fixes/14_07_26_fix5.md.")

    doc.core_properties.modified = datetime.now()
    doc.save(ARTICLE_PATH)
    print(ARTICLE_PATH)


if __name__ == "__main__":
    build_article_update()
