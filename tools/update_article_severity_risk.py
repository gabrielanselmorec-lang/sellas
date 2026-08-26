from __future__ import annotations

from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor

from update_article_abc import (
    GREEN,
    LIGHT_GREEN,
    add_bullet,
    add_caption,
    add_equation_box,
    set_cell_shading,
)


PROJECT_DIR = Path(__file__).resolve().parents[1]
ARTICLE_PATH = PROJECT_DIR / "docs" / "artigo_previsao_comportamental_bHave_integrado_mvp_logica.docx"
FIGURE_PATH = PROJECT_DIR / "docs" / "figures" / "figura_7_matriz_risco_c1_c2.png"
PART_TITLE = "PARTE V - Gravidade C1/C2, função informada e matriz de risco por local"


def draw_risk_matrix() -> None:
    FIGURE_PATH.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(10.8, 6.4), dpi=180)
    fig.patch.set_facecolor("#fffaf2")
    ax.set_facecolor("#fffaf2")
    quadrants = [
        (0, 0, "#dce8df", "Raro e pouco perigoso"),
        (0, 50, "#eedbd7", "Raro e muito perigoso"),
        (50, 0, "#e9d9b5", "Frequente e pouco perigoso"),
        (50, 50, "#dca9a3", "Frequente e muito perigoso"),
    ]
    for x, y, color, label in quadrants:
        ax.add_patch(plt.Rectangle((x, y), 50, 50, color=color, alpha=0.86, zorder=0))
        label_y = y + 30 if y == 50 else y + 42
        ax.text(x + 25, label_y, label, ha="center", va="center", fontsize=10.5, color="#51463b")

    points = [
        (18, 20, "o", "#587957", "C1: baixa ocorrência"),
        (78, 20, "o", "#587957", "C1: alta ocorrência"),
        (22, 100, "^", "#a94f49", "C2: baixa ocorrência"),
        (82, 100, "^", "#a94f49", "C2: alta ocorrência"),
        (50, 50, "x", "#6f6b66", "NC: editar classificação"),
    ]
    for x, y, marker, color, label in points:
        if marker == "x":
            ax.scatter(x, y, s=260, marker=marker, c=color, linewidths=2, zorder=3)
        else:
            ax.scatter(x, y, s=260, marker=marker, c=color, edgecolors="#fffdf8", linewidths=2, zorder=3)
        ax.annotate(label, (x, y), xytext=(0, -23), textcoords="offset points", ha="center", fontsize=9.5)

    ax.axvline(50, color="#8f8171", linestyle="--", linewidth=1)
    ax.axhline(50, color="#8f8171", linestyle="--", linewidth=1)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 105)
    ax.set_xlabel("Probabilidade de ocorrência da cadeia no local (%)", fontsize=11)
    ax.set_ylabel("Índice de perigo (%)", fontsize=11)
    ax.set_title("Matriz frequência × perigo por local", fontsize=15, fontweight="bold", color="#241a12", pad=14)
    ax.grid(color="#ffffff", linewidth=0.8, alpha=0.5)
    fig.text(
        0.5,
        0.01,
        "Círculo = C1 leve | Triângulo = C2 intenso | X = não classificado",
        ha="center",
        fontsize=10,
        color="#405f3d",
        fontweight="bold",
    )
    fig.tight_layout(rect=[0, 0.04, 1, 1])
    fig.savefig(FIGURE_PATH, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def update_article() -> None:
    if not ARTICLE_PATH.exists():
        raise FileNotFoundError(ARTICLE_PATH)
    draw_risk_matrix()
    document = Document(ARTICLE_PATH)
    if any(PART_TITLE in paragraph.text for paragraph in document.paragraphs):
        print(f"Parte já presente: {PART_TITLE}")
        return

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(PART_TITLE, level=1)
    lead = document.add_paragraph()
    lead_run = lead.add_run(
        "Esta parte documenta a classificação explícita da gravidade do comportamento interferente, "
        "o registro da função como hipótese informada, a edição auditável dos acontecimentos e a nova matriz "
        "de frequência × perigo filtrada por ambiente."
    )
    lead_run.bold = True

    document.add_heading("1. Classificação operacional C1 e C2", level=2)
    document.add_paragraph(
        "C1 e C2 são classes operacionais do acontecimento observado, não diagnósticos. C1 representa um episódio "
        "sem lesão, ferimento, sangramento ou direcionamento a ponto vital. C2 exige ao menos um desses critérios. "
        "A API rejeita combinações incoerentes, como C2 sem critério marcado ou C1 com sangramento."
    )
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Classe", "Critério", "Código numérico", "Índice visual de perigo"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, GREEN)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    rows = [
        ("C1 - leve", "Nenhum critério de dano presente", "1", "0,20"),
        ("C2 - intenso", "Lesão/ferimento, sangramento ou ponto vital", "2", "1,00"),
        ("NC", "Registro histórico ainda não editado", "nulo", "0,50 apenas para posição provisória"),
    ]
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cells[column_index].text = value
            if row_index % 2 == 0:
                set_cell_shading(cells[column_index], LIGHT_GREEN)

    document.add_heading("2. Função informada como hipótese", level=2)
    document.add_paragraph(
        "A função é armazenada junto ao comportamento como hipótese informada pelo profissional: atenção social, "
        "acesso a item ou atividade, fuga/esquiva, reforçamento automático, indeterminada/em avaliação ou texto livre. "
        "O campo não transforma associação em causalidade e pode ser corrigido posteriormente."
    )
    add_bullet(document, "Função informada não equivale a função demonstrada por avaliação funcional.")
    add_bullet(document, "Mudanças de função permanecem no log com o valor anterior e o valor atualizado.")

    document.add_heading("3. Matriz de risco por ambiente", level=2)
    document.add_paragraph(
        "O usuário seleciona um local e cada cadeia A-B-C aparece como um marcador. O eixo horizontal é a proporção "
        "da cadeia entre os registros daquele local. O eixo vertical é o índice de perigo calculado apenas com os "
        "registros classificados; registros antigos sem C1/C2 ficam no centro como pendência de edição."
    )
    add_equation_box(
        document,
        "Equações da matriz frequência × perigo",
        [
            "p(c|e) = n(c,e) / n(e)",
            "D(c,e) = [0,20·n_C1(c,e) + 1,00·n_C2(c,e)] / [n_C1(c,e) + n_C2(c,e)]",
            "R(c,e) = p(c|e) · D(c,e)",
        ],
    )
    document.add_paragraph(
        "Os coeficientes 0,20 e 1,00 organizam a exibição e não constituem uma escala clínica validada. O mapa serve "
        "para priorização visual e planejamento de observação, não para decidir intervenção automaticamente."
    )
    figure_paragraph = document.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(str(FIGURE_PATH), width=Inches(6.45))
    add_caption(document, "Figura 7. Matriz de frequência × perigo filtrada por local, com símbolos C1, C2 e não classificado.")

    document.add_heading("4. Gravidade e função na previsão", level=2)
    document.add_paragraph(
        "A previsão mantém a estimativa ponderada de P(B|A,E) e acrescenta a classificação S e a função F como "
        "fatores do registro-alvo. Dessa forma, o painel pode estimar a chance do comportamento em geral ou a chance "
        "conjunta de comportamento, nível e função selecionados."
    )
    add_equation_box(
        document,
        "Fatoração do registro-alvo",
        [
            "P(B,S,F|A,E) ≈ P_pond(B|A,E) × P(S|B,A,E) × P(F|B,A,E)",
            "p_suavizada = (k + 1) / (n + 2)",
            "P_pond(B|A,E) = Σ(w_j·p_j) / Σw_j",
        ],
    )
    document.add_paragraph(
        "A suavização Beta(1,1) reduz extremos em amostras pequenas. Se C1/C2 ou função não forem selecionados, "
        "o fator correspondente vale 1 e a previsão retorna à probabilidade do comportamento."
    )

    document.add_heading("5. Edição, auditoria e Excel", level=2)
    add_bullet(document, "Data, hora, segundo, ambiente, antecedente, comportamento e consequência podem ser corrigidos.")
    add_bullet(document, "C1/C2, critérios objetivos e função podem ser adicionados aos registros históricos.")
    add_bullet(document, "O log registra o estado anterior e o novo estado como ação registro_editado.")
    add_bullet(document, "O Excel passa a exibir classificação, função, critérios, perigo e valores anteriores da edição.")

    document.add_heading("6. Limites clínicos", level=2)
    document.add_paragraph(
        "Frequência alta não implica gravidade alta, e gravidade alta não implica frequência alta. A matriz mantém "
        "essas dimensões separadas para evitar que episódios raros e perigosos sejam ocultados por médias. Ainda assim, "
        "os resultados dependem da qualidade da observação, da oportunidade de resposta, da definição operacional e da "
        "consistência do preenchimento. As hipóteses funcionais exigem avaliação profissional apropriada."
    )

    document.add_heading("7. Verificação desta evolução", level=2)
    add_bullet(document, "Critérios C1/C2 e a fatoração da previsão foram cobertos por testes automatizados.")
    add_bullet(document, "Inserção e edição foram validadas no Supabase dentro de transação descartada por rollback.")
    add_bullet(document, "Filtro por local, matriz de risco e editor foram verificados no navegador sem traceback.")
    add_bullet(document, "Planilha de QA foi gerada com classificação, função, critérios e valores anteriores.")

    document.core_properties.modified = datetime.now()
    document.save(ARTICLE_PATH)
    print(ARTICLE_PATH)


if __name__ == "__main__":
    update_article()
